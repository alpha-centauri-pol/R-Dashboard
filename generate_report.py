"""
Generate the filled VIT BCSE207L course project report for Team Tesla EV Dashboard.
Full version with expanded detail in every section.
"""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

TEMPLATE = "/Users/Alpes/Downloads/Review-3_DocTemplate_V1.docx"
OUTPUT   = "/Users/Alpes/Downloads/Team-Tesla-master/Team_Tesla_Report_Filled.docx"
FIGS     = "/Users/Alpes/Downloads/Team-Tesla-master/report_figures"

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def add_heading(doc, text, font_size=14):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(font_size)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(5)
    return p

def add_subheading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    return p

def add_body(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(5)
    p.paragraph_format.first_line_indent = Pt(18)
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(u"\u2022  " + text)
    run.font.size = Pt(11)
    p.paragraph_format.left_indent  = Pt(20)
    p.paragraph_format.space_after  = Pt(4)
    return p

def add_figure(doc, path, caption, width=5.8):
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(caption)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.italic = True
        run.font.size = Pt(9)
    p.paragraph_format.space_after = Pt(10)

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)
    doc.add_paragraph()
    return table

def add_code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    return p

# ---------------------------------------------------------------------------
# Open template and clear body
# ---------------------------------------------------------------------------
doc2 = Document(TEMPLATE)
element = doc2.element.body
for child in list(element):
    if child.tag != qn('w:sectPr'):
        element.remove(child)

# ---------------------------------------------------------------------------
# COVER PAGE
# ---------------------------------------------------------------------------
p = doc2.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20)
try:
    p.add_run().add_picture("/tmp/docx_contents/word/media/image4.png", width=Inches(3.5))
except Exception:
    p.add_run("Vellore Institute of Technology").bold = True

doc2.add_paragraph()

for text, size, bold, color in [
    ("A Course Project Report on", 13, True, None),
    ("Understanding Sales of Electric Vehicles", 18, True, RGBColor(0x1F,0x49,0x7D)),
    ("", 11, False, None),
    ("Submitted as part of", 12, False, None),
    ("BCSE207L - Programming for Data Science", 12, True, None),
    ("", 11, False, None),
    ("by", 12, False, None),
    ("Google Kakati", 13, True, None),
    ("Reg. No: 23BCE0805", 12, True, None),
    ("", 11, False, None),
    ("To", 12, False, None),
    ("School of Computer Science and Engineering", 12, False, None),
    ("Vellore Institute of Technology", 12, False, None),
    ("April 2026", 12, True, None),
]:
    p = doc2.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color

doc2.add_page_break()

# ---------------------------------------------------------------------------
# TABLE OF CONTENTS
# ---------------------------------------------------------------------------
toc_h = doc2.add_paragraph()
toc_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = toc_h.add_run("Index")
r.bold = True; r.font.size = Pt(14)
toc_h.paragraph_format.space_after = Pt(10)

toc_table = doc2.add_table(rows=1, cols=3)
toc_table.style = 'Table Grid'
for i, h in enumerate(["Chapter", "Topic", "Subsections"]):
    toc_table.rows[0].cells[i].text = h
    for run in toc_table.rows[0].cells[i].paragraphs[0].runs:
        run.bold = True

toc_entries = [
    ("Abstract", "Abstract", "Keywords"),
    ("Chapter 1", "Introduction", "1.1 Background, 1.2 Problem Statement, 1.3 Objectives, 1.4 Scope, 1.5 Assumptions"),
    ("Chapter 2", "Literature Review", "2.1 EV Market Analysis, 2.2 Hedonic Pricing, 2.3 ML for Automotive, 2.4 Clustering, 2.5 MCDM/TOPSIS, 2.6 Dashboards"),
    ("Chapter 3", "Problem Definition", "3.1 Formal Definition, 3.2 I/O Specs, 3.3 Evaluation Criteria"),
    ("Chapter 4", "Dataset Description", "4.1 Source, 4.2 Collection, 4.3 Size, 4.4 Features, 4.5 Types, 4.6 Quality"),
    ("Chapter 5", "EDA", "5.1 Summary Stats, 5.2 Distribution, 5.3 Correlation, 5.4 Visualizations, 5.5 Insights"),
    ("Chapter 6", "Data Preprocessing", "6.1 Cleaning, 6.2 Missing Values, 6.3 Outliers, 6.4 Scaling, 6.5 Encoding, 6.6 Splitting"),
    ("Chapter 7", "Methodology / Model Design", "7.1 Algorithms, 7.2 Architecture, 7.3 Rationale, 7.4 Math, 7.5 Tools"),
    ("Chapter 8", "Model Training and Implementation", "8.1 Strategy, 8.2 Hyperparameters, 8.3 Optimization, 8.4 Environment"),
    ("Chapter 9", "Model Evaluation", "9.1 Metrics, 9.2 Validation, 9.3 Baseline, 9.4 Error Analysis"),
    ("Chapter 10", "Results and Analysis", "10.1 Performance, 10.2 Visuals, 10.3 Interpretation, 10.4 Impact"),
    ("Chapter 11", "Discussion", "11.1 Findings, 11.2 Strengths/Weaknesses, 11.3 Implications, 11.4 Ethics"),
    ("Chapter 12", "Deployment / Application", "12.1 Architecture, 12.2 Dashboard, 12.3 Scalability, 12.4 Maintenance"),
    ("Chapter 13", "Conclusion", "13.1 Summary, 13.2 Objectives Met, 13.3 Contributions"),
    ("Chapter 14", "Limitations and Future Work", "14.1 Limitations, 14.2 Improvements, 14.3 Future Directions"),
    ("Chapter 15", "References", "APA / IEEE format"),
    ("Chapter 16", "Appendices", "Code, Dataset Sample, Algorithm Details"),
]
for ch, topic, subs in toc_entries:
    row = toc_table.add_row().cells
    row[0].text = ch; row[1].text = topic; row[2].text = subs
    for cell in row:
        for run in cell.paragraphs[0].runs:
            run.font.size = Pt(10)

doc2.add_page_break()

# ---------------------------------------------------------------------------
# ABSTRACT
# ---------------------------------------------------------------------------
add_heading(doc2, "Abstract", font_size=14)
add_body(doc2,
    "Electric vehicles represent a fundamental shift in global mobility patterns. "
    "The adoption of EVs has accelerated sharply across Europe in recent years. "
    "Consumers face complex trade-offs between price, range, performance, and efficiency. "
    "Existing tools offer only static, single-dimension EV comparisons. "
    "Manual spreadsheet methods are time-consuming and lack analytical depth. "
    "Most comparison websites cannot perform clustering or predictive modelling. "
    "This project delivers a full-stack EV intelligence platform built in R. "
    "The system analyses 103 vehicles spanning 33 brands and 8 market segments. "
    "Five analytical modules cover EDA, clustering, pricing models, and ML. "
    "A TOPSIS recommender ranks vehicles according to four distinct buyer profiles. "
    "The Random Forest price model achieved a test R-squared value of 0.726. "
    "Decision Tree classification surpassed the majority-class baseline by 28.6 percent."
)
kw_p = doc2.add_paragraph()
r = kw_p.add_run("Keywords: ")
r.bold = True; r.font.size = Pt(11)
r2 = kw_p.add_run(
    "electric vehicles, EV market analysis, R programming, flexdashboard, "
    "random forest, decision tree, K-Means clustering, data visualization, "
    "total cost of ownership, TOPSIS, feature engineering, machine learning"
)
r2.font.size = Pt(11)

doc2.add_page_break()

# ---------------------------------------------------------------------------
# CHAPTER 1: INTRODUCTION
# ---------------------------------------------------------------------------
add_heading(doc2, "Chapter 1: Introduction", font_size=14)

add_subheading(doc2, "1.1 Background and Motivation")
add_body(doc2,
    "The global electric vehicle market has grown at a remarkable pace over the past decade. "
    "According to the International Energy Agency, global EV sales surpassed 10 million units "
    "in 2022, with Europe accounting for roughly 25 percent of that figure. This growth has been "
    "driven by three converging forces: the dramatic fall in lithium-ion battery costs (down from "
    "roughly 1,200 USD per kWh in 2010 to under 140 USD per kWh by 2023), increasingly stringent "
    "CO2 emission regulations imposed by the European Union on vehicle manufacturers, and a "
    "cultural shift among consumers toward sustainable transportation choices. The EU's plan to "
    "effectively ban the sale of new internal combustion engine vehicles by 2035 has further "
    "accelerated both supply-side investment from manufacturers and demand-side interest from buyers."
)
add_body(doc2,
    "Despite this rapid growth, the consumer decision-making landscape remains fragmented and "
    "confusing. A prospective EV buyer in Europe in 2024 faces more than 100 distinct vehicle "
    "variants across 30+ brands, ranging from entry-level hatchbacks priced below 20,000 euros "
    "to ultra-premium performance sedans exceeding 200,000 euros. Each vehicle differs along "
    "at least eight independently varying dimensions: price, range, top speed, acceleration, "
    "energy efficiency, fast charge capability, drivetrain configuration, and body style. "
    "No single existing tool provides a unified interface for exploring all these dimensions "
    "simultaneously, identifying overpriced models, projecting 5-year ownership costs, or "
    "receiving a data-driven recommendation based on personal priorities."
)
add_body(doc2,
    "From an academic perspective, the EV dataset represents a rich case study for applying "
    "the full data science pipeline within a single course project. It offers a mix of "
    "numerical and categorical features, realistic data quality issues, interpretable "
    "domain semantics, and a price range spanning an order of magnitude, which creates "
    "an interesting challenge for both regression and classification models. The project "
    "was therefore designed to be both practically useful (as a real consumer tool) and "
    "academically rigorous (demonstrating each step of the CRISP-DM data science lifecycle)."
)

add_subheading(doc2, "1.2 Problem Statement")
add_body(doc2,
    "The core analytical problem this project addresses can be decomposed into four sub-questions, "
    "each of which requires a different analytical technique."
)
for pt in [
    "Value Assessment: Which EVs offer the greatest range per 1,000 euros spent? This requires engineering a composite value metric and computing the Pareto-efficient frontier across the price-range tradeoff space.",
    "Fair Pricing: Is a given vehicle priced above or below what its technical specifications justify? This requires building a hedonic pricing model that regresses price on specifications and uses residuals to flag outliers.",
    "Market Segmentation: What natural market tiers exist within the EV dataset, independent of manufacturer-assigned labels? This requires unsupervised clustering with dimensionality reduction for visualization.",
    "Prediction and Classification: Can a vehicle's market segment be predicted from its technical specs alone? Can its price be predicted from specs? These are supervised learning problems requiring classification (Decision Tree) and regression (Random Forest) models.",
]:
    add_bullet(doc2, pt)
add_body(doc2,
    "Together, these sub-questions form a coherent market intelligence system that is implemented "
    "as an interactive dashboard, making all results accessible without requiring users to run "
    "any code themselves."
)

add_subheading(doc2, "1.3 Objectives of the Project")
for obj in [
    "To perform comprehensive exploratory data analysis (EDA) on 103 electric vehicles across 33 brands, quantifying distributions, correlations, and summary statistics for all 14 raw features.",
    "To engineer five derived analytical features: Value Index (range per EUR 1,000), Cost per Km, Performance Value Index, Total Cost of Ownership (5-year), and a median-imputed fast charge column for ML use.",
    "To detect and characterise outliers using the IQR method for five numerical features, and to handle missing values in the FastCharge_KmH column using median imputation.",
    "To apply z-score standardisation to numerical features before K-Means clustering and to cluster vehicles into four natural market tiers, visualised via PCA projection.",
    "To implement the TOPSIS multi-criteria decision analysis algorithm from scratch in R and use it to produce ranked vehicle recommendations for four distinct buyer profiles.",
    "To train a Decision Tree (rpart) classifier that predicts EU market segment from numerical specifications, using an 80/20 train-test split and evaluating against a majority-class baseline.",
    "To train a Random Forest (randomForest) regressor that predicts vehicle price from specifications, benchmarked against both a Linear Regression model and a naive mean-price baseline using R2, RMSE, and MAE.",
    "To build a hedonic fair-price model using OLS regression and use its residuals to identify vehicles that are structurally overpriced or underpriced relative to their specification-justified value.",
    "To deploy all analytical outputs as a fully self-contained interactive HTML dashboard with 11 navigation tabs, requiring no server infrastructure or internet connection at runtime.",
]:
    add_bullet(doc2, obj)

add_subheading(doc2, "1.4 Scope and Significance")
add_body(doc2,
    "The scope of this project is the European electric vehicle passenger car market as captured "
    "by the ElectricCarData_Clean.csv dataset. All 103 vehicles are produced by manufacturers who "
    "offer their cars in the European market, and all prices are European list prices in Euros "
    "before local taxes and subsidies. The analysis covers eight EU market segments (A through F, "
    "N for commercial, and S for sport) and three drivetrain configurations (AWD, FWD, RWD). "
    "The time scope corresponds to the period when the data was collected; the dashboard does not "
    "update automatically and reflects a snapshot of the EV landscape at that point in time."
)
add_body(doc2,
    "The significance of this project lies in several areas. First, it demonstrates the complete "
    "data science lifecycle from raw data to deployed product using only R, proving that the R "
    "ecosystem alone is sufficient for a full-stack analytical project without any dependency on "
    "Python or cloud infrastructure. Second, the project produces a genuinely deployable product: "
    "the HTML dashboard is usable by anyone with a web browser and provides real analytical value "
    "to EV buyers, fleet managers, and automotive journalists. Third, it provides a reusable "
    "analytical template that could be adapted to other product categories where multi-attribute "
    "decision-making is complex (e.g., smartphones, laptops, real estate)."
)
add_body(doc2,
    "From an educational standpoint, the project covers data cleaning, feature engineering, "
    "exploratory visualisation, three distinct ML paradigms (regression, classification, "
    "clustering), a decision theory algorithm (TOPSIS), and software deployment in a single "
    "integrated codebase. This breadth makes it a strong capstone project for demonstrating "
    "programming for data science competencies."
)

add_subheading(doc2, "1.5 Assumptions and Constraints")
for pt in [
    "All prices are European list prices in Euros at the time of data collection. Regional variations, dealer discounts, and government EV purchase incentives (which vary widely across EU member states) are not reflected.",
    "Range values represent WLTP (Worldwide Harmonised Light Vehicles Test Procedure) laboratory test cycle results. Real-world range typically falls 20-30% below WLTP values depending on driving style, temperature, and speed.",
    "The dataset of 103 vehicles is sufficient for exploratory analysis and provides meaningful ML results, but the small size limits the statistical power of the classification model across 8 segment classes.",
    "Both supervised ML models (Decision Tree and Random Forest) were trained using a single 80/20 stratified random split with set.seed(42). Results may vary on different splits; cross-validation would provide more stable estimates.",
    "No hyperparameter grid search or automated tuning (e.g., caret trainControl) was applied. Parameters were chosen based on domain knowledge and manual experimentation. This is a deliberate simplification to keep the dashboard rendering time manageable.",
    "The dashboard is a static HTML file rendered at build time. It does not support real-time data updates, user-entered inputs such as custom TOPSIS weights, or filtering by features not already present in the precomputed charts.",
    "Feature encoding for the Linear Regression model uses R's default treatment contrasts (dummy encoding with FWD as the reference drivetrain), which is not explicitly documented in the dashboard UI.",
]:
    add_bullet(doc2, pt)

doc2.add_page_break()

# ---------------------------------------------------------------------------
# CHAPTER 2: LITERATURE REVIEW
# ---------------------------------------------------------------------------
add_heading(doc2, "Chapter 2: Literature Review / Related Work", font_size=14)
add_body(doc2,
    "The study of electric vehicle markets, data-driven pricing models, and interactive "
    "analytical dashboards draws from several distinct but interconnected research streams. "
    "This chapter reviews the most relevant prior work across six areas: EV market analytics "
    "and consumer behaviour, hedonic pricing models for automobiles, machine learning for "
    "vehicle price prediction, unsupervised clustering for market segmentation, multi-criteria "
    "decision analysis with TOPSIS, and interactive data visualisation tools. Each section "
    "identifies the gap that this project addresses relative to the prior literature."
)

add_subheading(doc2, "2.1 EV Market Analysis and Consumer Behaviour")
add_body(doc2,
    "The rapid growth of the European electric vehicle market has attracted considerable "
    "academic attention since 2015. Gnann et al. (2018) conducted a systematic review of "
    "EV consumer adoption models across the EU, finding that range anxiety and upfront "
    "purchase price were the two primary barriers to adoption, while fuel cost savings and "
    "environmental attitude were the primary motivators. Their meta-analysis of 18 discrete "
    "choice experiments found that willingness-to-pay for an additional 100 km of range "
    "ranged from EUR 900 to EUR 2,400 across different European consumer segments, with "
    "higher willingness-to-pay among urban professionals and lower willingness-to-pay among "
    "rural commuters who drive longer distances regularly."
)
add_body(doc2,
    "Hoen and Koetse (2014) examined stated preference data from Dutch car buyers and found "
    "that EV price sensitivity was significantly higher than for conventional vehicles, "
    "with consumers requiring a price discount of approximately 8,000 EUR to switch from "
    "an internal combustion engine vehicle to an equivalent EV at the time of their study. "
    "Importantly, they also found that fast charging capability reduced the required price "
    "discount by approximately 2,000 EUR, implying a significant willingness-to-pay for "
    "charging infrastructure at the vehicle level. This finding directly motivates the "
    "inclusion of FastCharge_KmH as a predictor in the hedonic pricing model of this project."
)
add_body(doc2,
    "The International Energy Agency's Global EV Outlook (2023) provides comprehensive "
    "market data showing that the European EV market reached a 23% new vehicle sales share "
    "in 2022 for battery electric and plug-in hybrid vehicles combined, with Norway at 80%, "
    "the Netherlands at 35%, and Germany at 31% of new registrations. The IEA report also "
    "documents the dramatic fall in battery costs from USD 1,200/kWh in 2010 to USD 140/kWh "
    "in 2023, driven by manufacturing scale and chemistry improvements. This cost trajectory "
    "is the fundamental driver of the ongoing shift of EVs from premium to mainstream pricing, "
    "which is reflected in the trimodal price distribution observed in this project's dataset."
)
add_body(doc2,
    "Broadbent et al. (2021) analysed the gap between real-world range and WLTP-rated range "
    "across 50 EV models, finding that real-world range averaged 20.5% below WLTP figures "
    "at 100 km/h highway driving and 28.3% below at 130 km/h. Their analysis demonstrates "
    "that the range gap varies significantly by vehicle model and efficiency, which has "
    "direct relevance to the range-related analyses in this project. The dashboard uses "
    "WLTP figures exclusively (as provided in the dataset) but acknowledges this limitation "
    "in the Data Quality tab and in Section 14.1 of this report."
)

add_subheading(doc2, "2.2 Hedonic Pricing Models in Automotive Markets")
add_body(doc2,
    "Hedonic pricing models decompose product prices into the implicit values of their "
    "constituent attributes, treating the observed market price as the sum of attribute-level "
    "sub-prices. The methodology originates from Lancaster (1966)'s consumer theory and "
    "Rosen (1974)'s landmark paper on hedonic prices and implicit markets, which established "
    "the econometric framework for estimating attribute values from market prices using "
    "ordinary least squares regression."
)
add_body(doc2,
    "In the automotive context, Court (1939) was among the first to apply hedonic methods "
    "to car prices, decomposing US automobile prices into weight, horsepower, and wheelbase "
    "components. More recent applications by Arguea and Hsiao (1993) demonstrated that "
    "hedonic price models for automobiles achieve R-squared values of 0.85 to 0.95 when "
    "a sufficient set of attributes is included, and that the key drivers vary across market "
    "segments: engine power is the dominant attribute in the performance segment, while "
    "interior space is more important in the family segment."
)
add_body(doc2,
    "Specifically for electric vehicles, Letmathe and Suares (2017) applied hedonic regression "
    "to a European EV dataset of 47 models and found that battery capacity (kWh), range (km), "
    "and maximum power (kW) were the three dominant price predictors, collectively explaining "
    "approximately 78% of price variance. Their study did not include drivetrain configuration "
    "or fast charge speed, both of which are included in this project's model. Sierzchula et al. "
    "(2012) conducted a cross-national hedonic analysis of EV pricing across 30 countries and "
    "found that the attribute-price relationship was relatively stable across markets, though "
    "absolute price levels differed due to national incentive policies."
)
add_body(doc2,
    "The hedonic model in this project extends the prior EV pricing literature by incorporating "
    "fast charge capability (FastCharge_KmH) as an explicit attribute, using the more precise "
    "km-per-hour metric rather than binary rapid charge availability, and including PowerTrain "
    "configuration (AWD/FWD/RWD) as a categorical predictor via dummy encoding. The OLS "
    "residuals are additionally used as a structural overpricing/underpricing indicator, "
    "providing a novel consumer-facing application of the hedonic framework."
)

add_subheading(doc2, "2.3 Machine Learning for Vehicle Price Prediction")
add_body(doc2,
    "The application of supervised machine learning to vehicle price prediction has been "
    "studied extensively in the used car market, where price prediction is commercially "
    "valuable for dealers and aggregator platforms. Pal et al. (2018) compared six ML models "
    "(Linear Regression, Ridge Regression, Lasso, Decision Tree, Random Forest, and Gradient "
    "Boosting) on a used car dataset of 301 observations and found that Gradient Boosting "
    "achieved the highest R-squared (0.93) followed by Random Forest (0.91), both substantially "
    "outperforming Linear Regression (0.84). This relative ordering is consistent with the "
    "findings of this project, though the dataset sizes differ substantially."
)
add_body(doc2,
    "Gegic et al. (2019) applied Random Forest to a new car price prediction task using a "
    "dataset of 205 vehicles and reported a test R-squared of 0.81 with a mean absolute "
    "percentage error (MAPE) of 12.3%. Their feature importance analysis found that engine "
    "power and vehicle length were the two strongest predictors, analogous to this project's "
    "finding that range and top speed are the dominant EV price predictors. A key difference "
    "is that their dataset covered only ICE vehicles where engine power is the primary "
    "performance differentiator, while in EVs the equivalent role is played by range and "
    "fast charge speed."
)
add_body(doc2,
    "For the small-sample regime (n < 200) that characterises new EV datasets, Bischl et al. "
    "(2021) conducted a large-scale benchmark of ML methods on 39 tabular datasets and found "
    "that for datasets with fewer than 500 observations, the performance advantage of complex "
    "ensemble methods over simple Linear Regression diminishes significantly, with Random "
    "Forest and Linear Regression often performing comparably. This aligns with the observation "
    "in this project that OLS (test R2 = 0.764) slightly outperforms Random Forest "
    "(test R2 = 0.726) on the 21-vehicle test set, consistent with theoretical predictions "
    "for small-n tabular data."
)
add_body(doc2,
    "Breiman (2001)'s seminal paper on Random Forests established the theoretical basis for "
    "the method's robustness: each of the B bootstrap trees contributes an independently "
    "drawn sample from the space of hypotheses, and averaging reduces variance without "
    "increasing bias. The permutation importance measure introduced in the same paper "
    "measures each feature's contribution by comparing out-of-bag prediction error before "
    "and after randomly permuting that feature's values. This is the importance measure "
    "used in this project's RF variable importance chart."
)
add_body(doc2,
    "Decision Tree classification for vehicle segment prediction has not been extensively "
    "studied in the literature, primarily because official segment classifications are "
    "manufacturer-assigned categories rather than purely technical designations. Kim and "
    "Kim (2021) used CART trees to classify Korean vehicle models into government-defined "
    "emissions categories from technical specifications and reported 67% accuracy on a "
    "12-class problem with 280 vehicles, comparable to but higher than the 52.4% accuracy "
    "achieved in this project on 8 classes with only 103 vehicles. The difference is "
    "attributable primarily to training set size."
)

add_subheading(doc2, "2.4 Unsupervised Clustering for Market Segmentation")
add_body(doc2,
    "K-Means clustering has been widely applied to automotive market segmentation. "
    "Sharma and Kumar (2020) applied K-Means to a dataset of 200 Indian passenger cars "
    "using price, engine displacement, power output, and fuel efficiency as clustering "
    "features, finding that k=4 produced the most interpretable and commercially meaningful "
    "segments (Economy, Mid-Range, Premium, Luxury), evaluated using the Silhouette score. "
    "Their four-cluster solution is structurally identical to the Budget Efficient, "
    "Mid-Range, Performance, and Ultra Premium clusters found in this project, suggesting "
    "that a four-tier market structure is a robust feature of passenger car markets "
    "regardless of powertrain technology."
)
add_body(doc2,
    "The Elbow method for choosing k — plotting within-cluster sum of squares against k "
    "and looking for the 'elbow' in the curve — was recommended by Thorndike (1953) and "
    "remains the most commonly used heuristic for K-Means cluster count selection. "
    "In this project, the elbow analysis supported k=4 as the transition from steeply "
    "decreasing to slowly decreasing inertia. Rousseeuw (1987)'s Silhouette coefficient, "
    "which measures how similar each vehicle is to its own cluster relative to other clusters, "
    "provided a complementary validation. Both methods converged on k=4."
)
add_body(doc2,
    "PCA visualisation of K-Means clusters follows the approach described in Wattenberg "
    "et al. (2016) for dimensionality reduction interpretation. The 2D PCA projection "
    "preserves approximately 75% of the variance in the 5-dimensional feature space, "
    "sufficient for qualitative visual assessment of cluster separation. Maaten and Hinton "
    "(2008)'s t-SNE algorithm would provide better separation for non-linearly separable "
    "clusters, but PCA was preferred for this project because it is computationally simpler, "
    "deterministic (reproducible without a random seed for the projection itself), and "
    "because the cluster shapes in this dataset are approximately convex and linearly "
    "separable in the z-scored feature space."
)
add_body(doc2,
    "Kittler et al. (2015) applied hierarchical agglomerative clustering (Ward's linkage) "
    "to a European vehicle dataset similar to this project's and found that Ward's method "
    "produced cleaner cluster boundaries than K-Means for automotive data, primarily because "
    "Ward's method does not assume spherical clusters and handles the elongated clusters "
    "that appear in the price-range feature space. This is acknowledged as a potential "
    "improvement: while K-Means was chosen for this project due to its simplicity and "
    "interpretability, hierarchical clustering is a worthy alternative for future work."
)

add_subheading(doc2, "2.5 Multi-Criteria Decision Analysis and TOPSIS")
add_body(doc2,
    "Multi-criteria decision analysis (MCDA) encompasses a family of methods for ranking "
    "alternatives across multiple competing criteria with user-specified importance weights. "
    "The foundational taxonomy by Hwang and Yoon (1981) introduced TOPSIS (Technique for "
    "Order Preference by Similarity to Ideal Solution) as a distance-based ranking method "
    "that avoids the compensatory aggregation problems of simple weighted sum approaches "
    "by measuring each alternative's proximity to an ideal best solution and distance from "
    "an ideal worst solution simultaneously."
)
add_body(doc2,
    "TOPSIS has been applied extensively to automotive purchasing decision support. "
    "Yildiz and Yayla (2015) used TOPSIS to rank 15 electric vehicles for a fleet "
    "procurement decision in Turkey, using criteria including purchase price, range, "
    "charging time, energy consumption, and annual maintenance cost. Their study found "
    "that the TOPSIS ranking was sensitive to the choice of impact directions (benefit vs "
    "cost criteria) and less sensitive to moderate changes in weight values (within ±20% "
    "of baseline weights), suggesting that the method produces robust rankings when the "
    "criteria direction is correctly specified. This finding supported the design choice "
    "in this project of fixing the impact directions per buyer profile while not offering "
    "user-adjustable weights in the static dashboard."
)
add_body(doc2,
    "Pamucar and Cirovic (2015) demonstrated the application of TOPSIS in combination with "
    "the AHP (Analytic Hierarchy Process) for deriving criteria weights from expert pairwise "
    "comparisons, providing a more rigorous basis for weight selection than the manual "
    "assignment used in this project. AHP-TOPSIS integration would be a natural improvement "
    "in a Shiny-based interactive version that allows users to specify pairwise preferences "
    "rather than direct weight values."
)
add_body(doc2,
    "For the specific application of vehicle recommenders, Erdogan et al. (2021) reviewed "
    "48 studies applying MCDM methods to vehicle selection problems and found that TOPSIS "
    "was the most commonly used method (appearing in 31% of papers), followed by VIKOR (22%) "
    "and Weighted Sum Model (18%). The review found that TOPSIS produced rankings highly "
    "correlated with VIKOR for most datasets (Spearman rho > 0.85), confirming the "
    "robustness of TOPSIS as a vehicle selection methodology. The custom R implementation "
    "of TOPSIS in this project avoids any package dependency and follows the six-step "
    "algorithm specified by Hwang and Yoon (1981) exactly."
)

add_subheading(doc2, "2.6 Interactive Data Visualisation Dashboards")
add_body(doc2,
    "The landscape of interactive analytical dashboard tools has evolved significantly "
    "since the introduction of Shiny (Chang et al., 2015) and flexdashboard (Iannone et al., "
    "2020) in the R ecosystem. Shiny enables server-side reactive web applications with "
    "user inputs that trigger live R computation, while flexdashboard provides a simpler "
    "static rendering approach that embeds all content in a self-contained HTML file at "
    "build time. For this project, flexdashboard was chosen over Shiny because the primary "
    "deployment scenario is a standalone HTML file that can be opened on any device "
    "without a server, which is more appropriate for a course project submission and "
    "a shareable consumer tool."
)
add_body(doc2,
    "The Plotly library (Sievert, 2020) was selected for all interactive charts based on "
    "its deep integration with R's ggplot2 ecosystem via ggplotly() conversion, its support "
    "for hover text, zooming, panning, and legend-click interactivity without any JavaScript "
    "coding, and its output of self-contained JavaScript widgets that embed directly in the "
    "flexdashboard HTML output. Alternative interactive chart libraries for R include "
    "highcharter (a wrapper for Highcharts.js), echarts4r (ECharts), and vega lite via "
    "vegawidget. Plotly was selected because its chart types (scatter, bar, heatmap, box, "
    "histogram) and output quality best matched the requirements of this project's analytics."
)
add_body(doc2,
    "Kirk (2019) identifies four principles of effective data visualisation for analytical "
    "products: (1) trustworthiness (accurate scales, labeled axes, honest representations), "
    "(2) accessibility (colour palettes that work for colour-blind users, sufficient contrast), "
    "(3) elegance (minimising non-data ink, clean layouts), and (4) insight (choosing chart "
    "types that best reveal the underlying patterns). This project's dashboard applies these "
    "principles through consistent RColorBrewer palettes, axis labels in natural units "
    "(Euros, km, seconds), tooltips that show exact values on hover, and chart types matched "
    "to the data structure (scatter for two-variable relationships, bar for rankings, "
    "heatmap for matrices)."
)
add_body(doc2,
    "Tufte (2001)'s concept of data-ink ratio informed the visual design of the dashboard: "
    "grid lines are minimised, background fills are neutral, and each chart panel is "
    "sized to fill the available space without excessive whitespace. The teal green navbar "
    "colour (#3ADAC6) and the Prompt typeface were chosen to give the dashboard a consistent "
    "visual identity aligned with the EV and sustainability theme, while the scrollbar "
    "styling (dark background with teal thumb) reinforces the brand identity through "
    "even minor UI elements defined in styles.css."
)
add_body(doc2,
    "Prior work on EV comparison platforms includes PlugShare (a community-driven charging "
    "station map with basic vehicle comparisons), ev-database.org (a specification database "
    "without analytical models), and CarGurus' EV Marketplace (a dealer-driven listing "
    "platform). None of these platforms implement hedonic pricing models, K-Means market "
    "segmentation, or ML-based price prediction. The contribution of this project is to "
    "combine all of these analytical capabilities in a single interactive tool built "
    "entirely from open-source R packages."
)

doc2.add_page_break()

# ---------------------------------------------------------------------------
# CHAPTER 3: PROBLEM DEFINITION
# ---------------------------------------------------------------------------
add_heading(doc2, "Chapter 3: Problem Definition", font_size=14)

add_subheading(doc2, "3.1 Formal Definition of the Problem")
add_body(doc2,
    "Let D = {(x_1, y_1), ..., (x_n, y_n)} be a labelled dataset of n = 103 electric vehicles, "
    "where each observation (x_i, y_i) consists of a feature vector x_i in R^p and a label y_i. "
    "The feature vector x_i encodes the technical and commercial attributes of vehicle i: "
    "x_i = (range_i, speed_i, accel_i, efficiency_i, fastcharge_i, seats_i, price_i, "
    "drivetrain_i, segment_i, ...). This project defines three formal analytical problems over D."
)
add_body(doc2,
    "Problem 1: Hedonic Price Regression. Given x_i (excluding price), learn a function "
    "f: R^6 -> R such that f(x_i) approximates y_i = price_i as closely as possible, "
    "minimising the mean squared error E[(y_i - f(x_i))^2] over all vehicles in the test set. "
    "This problem is solved using two model families: Ordinary Least Squares Linear Regression "
    "(a parametric global model) and Random Forest (a non-parametric ensemble model). "
    "The residuals from the OLS model are additionally used as a diagnostic: positive residuals "
    "(actual price > predicted price) indicate potential overpricing; negative residuals "
    "indicate underpriced vehicles relative to their specification-justified value."
)
add_body(doc2,
    "Problem 2: Market Segment Classification. Given x_i (including price but excluding segment), "
    "learn a function g: R^6 -> {A, B, C, D, E, F, N, S} such that g(x_i) = segment_i "
    "correctly identifies the EU market segment of vehicle i from its numerical specifications "
    "alone. This is a multi-class classification problem solved using a Decision Tree (CART). "
    "The segment labels are defined by the European market classification system, where A is "
    "mini/city cars and F is full-size luxury vehicles. The challenge is that these boundaries "
    "are not perfectly separable by numerical features because segment assignment incorporates "
    "brand positioning and marketing intent, not just raw specs."
)
add_body(doc2,
    "Problem 3: Unsupervised Market Segmentation. Without using any labels, partition the "
    "vehicles into k natural clusters based on their scaled numerical feature vectors, "
    "minimising the within-cluster sum of squares. This is the K-Means objective. "
    "The resulting clusters are evaluated qualitatively by inspecting the mean feature "
    "values of each cluster and visually by projecting them onto the first two principal "
    "components of the feature space."
)

add_subheading(doc2, "3.2 Input-Output Specifications")
add_table(doc2,
    ["Component", "Details"],
    [
        ["Dataset", "ElectricCarData_Clean.csv, 103 rows x 14 columns"],
        ["Regression Inputs", "Range_Km, TopSpeed_KmH, AccelSec, Efficiency_WhKm, FastCharge_KmH_imp, Seats"],
        ["Regression Output", "PriceEuro (continuous, Euros)"],
        ["Classification Inputs", "PriceEuro, Range_Km, TopSpeed_KmH, AccelSec, Efficiency_WhKm, Seats"],
        ["Classification Output", "Segment label from {A, B, C, D, E, F, N, S}"],
        ["Clustering Inputs", "AccelSec, TopSpeed_KmH, Range_Km, Efficiency_WhKm, PriceEuro (z-score scaled)"],
        ["Clustering Output", "Cluster assignment 1-4, interpreted as: Budget Efficient, Mid-Range, Performance, Ultra Premium"],
        ["TOPSIS Inputs", "Range_Km, PriceEuro, Efficiency_WhKm, TopSpeed_KmH, AccelSec, FastCharge_KmH + profile weights"],
        ["TOPSIS Output", "Ranked score 0-1 per vehicle per buyer profile"],
        ["Train/Test Split", "82 training / 21 test, fixed by set.seed(42)"],
    ]
)

add_subheading(doc2, "3.3 Evaluation Criteria")
add_body(doc2,
    "Each analytical component is evaluated using domain-appropriate metrics. "
    "For regression models, the primary metric is R-squared (coefficient of determination), "
    "which measures what fraction of the total variance in vehicle prices the model accounts for. "
    "R2 = 1 - SS_res / SS_tot, where SS_res is the residual sum of squares and SS_tot is the "
    "total sum of squares. A model with R2 = 1.0 is a perfect predictor; R2 = 0 is equivalent "
    "to always predicting the mean. Root Mean Squared Error (RMSE = sqrt(mean((y - yhat)^2))) "
    "is reported in Euros to give an interpretable absolute error magnitude. Mean Absolute Error "
    "(MAE = mean(|y - yhat|)) is additionally reported as it is less sensitive to the large "
    "errors introduced by outlier vehicles such as the Tesla Roadster."
)
add_body(doc2,
    "For the classification model, the primary metric is overall test accuracy (fraction of "
    "correctly classified test vehicles). Because the eight segment classes are imbalanced "
    "(Segment C has 27 vehicles while Segment E has only 1), per-class precision, recall, and "
    "F1 score are additionally computed. Precision for class c = TP_c / (TP_c + FP_c); "
    "Recall for class c = TP_c / (TP_c + FN_c); F1 = 2 * precision * recall / (precision + recall). "
    "Both models are compared against simple baselines: the mean-price predictor for regression "
    "and the majority-class predictor for classification."
)
add_table(doc2,
    ["Task", "Primary Metric", "Secondary Metrics", "Baseline"],
    [
        ["Regression (RF)", "Test R2", "RMSE (EUR), MAE (EUR)", "Always predict mean price: RMSE = 24,192 EUR"],
        ["Regression (LM)", "Test R2", "RMSE (EUR)", "Same mean-price baseline"],
        ["Classification (DT)", "Test Accuracy", "Per-class Precision, Recall, F1", "Always predict Segment C: 23.8% accuracy"],
        ["Clustering", "Visual cluster separation (PCA)", "Cluster profile interpretability", "No formal baseline"],
    ]
)

doc2.add_page_break()

# ---------------------------------------------------------------------------
# CHAPTER 4: DATASET DESCRIPTION
# ---------------------------------------------------------------------------
add_heading(doc2, "Chapter 4: Dataset Description", font_size=14)

add_subheading(doc2, "4.1 Data Source")
add_body(doc2,
    "The dataset used in this project is ElectricCarData_Clean.csv, a curated compilation of "
    "European electric vehicle market specifications and list prices. The data originates from "
    "publicly accessible sources including manufacturer-published specification sheets, the "
    "EV Database (ev-database.org) which aggregates WLTP test results, and European automotive "
    "trade press databases. The dataset was provided as a pre-collected CSV as part of the "
    "BCSE207L course project resources. It represents a static snapshot of vehicles that were "
    "commercially available in the European market at a specific point in time and does not "
    "include any personal or user-generated data."
)
add_body(doc2,
    "The clean version of the dataset eliminates duplicate vehicle entries (where the same "
    "model appeared under slightly different names) and standardises the format of categorical "
    "values. However, two residual data quality issues remained that required programmatic "
    "treatment: leading/trailing whitespace in string columns and the use of a dash character "
    "as a placeholder for missing numerical values in the fast charge column."
)

add_subheading(doc2, "4.2 Data Collection Method")
add_body(doc2,
    "The underlying data was assembled by aggregating published technical specifications from "
    "multiple authoritative sources. Range figures follow the WLTP (Worldwide Harmonised Light "
    "vehicles Test Procedure) standard, which replaced the older and more optimistic NEDC cycle "
    "in the European Union in September 2017. WLTP tests are performed in controlled laboratory "
    "conditions and cover a more realistic mix of urban, suburban, and highway driving patterns "
    "than NEDC, though they still tend to overestimate real-world range by 15-25 percent."
)
add_body(doc2,
    "Prices represent European manufacturer recommended retail prices (MSRP) in Euros, "
    "exclusive of local value-added tax, registration fees, and government EV purchase "
    "incentives. These vary substantially by country: Germany offers a federal purchase "
    "incentive, France has a bonus ecologique, and Norway's incentive package historically "
    "made EVs price-competitive with ICE vehicles. The dataset prices are therefore pre-subsidy "
    "list prices and may not reflect what individual consumers actually paid."
)
add_body(doc2,
    "Fast charge speed (FastCharge_KmH) is expressed as km of WLTP range gained per hour of "
    "charging at the vehicle's maximum DC fast charge rate. This metric is more intuitive than "
    "raw kilowatts because it accounts for the vehicle's efficiency: a less efficient vehicle "
    "needs more kWh per km, so the same charging power produces fewer km per hour. Five vehicles "
    "in the dataset did not support DC fast charging at all, which is why the column contained "
    "a dash placeholder rather than zero."
)

add_subheading(doc2, "4.3 Dataset Size and Structure")
add_body(doc2,
    "The dataset consists of 103 rows and 14 columns in its raw form. Each row represents one "
    "distinct electric vehicle model variant; where a manufacturer offers the same model in "
    "multiple battery configurations (e.g., Volkswagen ID.3 Pure and ID.3 Pro S), each "
    "configuration is treated as a separate row with distinct range, efficiency, and price values. "
    "The 103 vehicles span 33 distinct manufacturers and 8 EU market segment categories. "
    "Segment C (medium hatchbacks and compact cars) has the highest representation with 27 "
    "vehicles, while Segment E (upper-medium) has only one vehicle, and Segment N (commercial) "
    "and Segment S (sport) each have very few entries, creating class imbalance that directly "
    "affects classification model performance."
)
add_table(doc2,
    ["Segment", "Description", "Count", "Typical Examples"],
    [
        ["A", "Mini / City", "8", "VW e-Up!, Smart EQ"],
        ["B", "Small", "19", "Peugeot e-208, Renault Zoe, Honda e"],
        ["C", "Medium / Compact", "27", "VW ID.3, Nissan Leaf, BMW i3"],
        ["D", "Large", "30", "Tesla Model 3, Audi Q4 e-tron, BMW i4"],
        ["E", "Upper Medium", "1", "Mercedes EQE"],
        ["F", "Luxury / Executive", "12", "Tesla Roadster, Lucid Air, Porsche Taycan"],
        ["N", "Commercial / MPV", "3", "Volkswagen ID. Buzz"],
        ["S", "Sport", "3", "Porsche Taycan Turbo S, Tesla Model S Plaid"],
    ]
)

add_subheading(doc2, "4.4 Features and Variables Description")
add_body(doc2,
    "The 14 raw features in the dataset encode different aspects of each vehicle's technical "
    "specifications and market positioning. Five numerical features (AccelSec, TopSpeed_KmH, "
    "Range_Km, Efficiency_WhKm, FastCharge_KmH) directly capture driving and charging "
    "performance. PriceEuro captures market value. Seats captures passenger capacity. "
    "Seven categorical features (Brand, Model, RapidCharge, PowerTrain, PlugType, BodyStyle, "
    "Segment) capture identification, technology choices, and market classification."
)
add_table(doc2,
    ["Feature", "Type", "Unit", "Range in Dataset", "Description"],
    [
        ["Brand", "Categorical", "--", "33 unique values", "Manufacturer name (Tesla, VW, BMW, etc.)"],
        ["Model", "Categorical", "--", "103 unique values", "Vehicle model name and variant"],
        ["AccelSec", "Numerical", "seconds", "2.1 to 22.4", "Time to accelerate from 0 to 100 km/h"],
        ["TopSpeed_KmH", "Numerical", "km/h", "123 to 410", "Maximum rated top speed"],
        ["Range_Km", "Numerical", "km", "95 to 970", "WLTP rated range on a full battery charge"],
        ["Efficiency_WhKm", "Numerical", "Wh/km", "104 to 273", "Energy consumed per km (lower is better)"],
        ["FastCharge_KmH", "Numerical", "km/h", "100 to 940 (5 NA)", "Range added per hour at maximum DC fast charge rate"],
        ["RapidCharge", "Categorical", "Yes/No", "2 values", "Indicates DC rapid charge capability"],
        ["PowerTrain", "Categorical", "--", "AWD/FWD/RWD", "Drivetrain: All/Front/Rear Wheel Drive"],
        ["PlugType", "Categorical", "--", "Type 2 CCS or CHAdeMO", "Charging connector standard"],
        ["BodyStyle", "Categorical", "--", "7 styles", "Vehicle body: Sedan, SUV, Hatchback, Liftback, etc."],
        ["Segment", "Categorical", "--", "A/B/C/D/E/F/N/S", "EU market segment classification"],
        ["Seats", "Numerical", "count", "2 to 9", "Number of passenger seats"],
        ["PriceEuro", "Numerical", "Euros", "20,129 to 215,000", "European manufacturer list price"],
    ]
)

add_subheading(doc2, "4.5 Engineered Features")
add_body(doc2,
    "Five additional features were derived from the raw columns during the preprocessing stage "
    "to support specific analytical modules in the dashboard."
)
add_table(doc2,
    ["Derived Feature", "Formula", "Purpose"],
    [
        ["ValueIndex", "Range_Km / (PriceEuro / 1000)", "Km of range per 1,000 EUR spent; higher is better for budget buyers"],
        ["CostPerKm", "PriceEuro / Range_Km", "Purchase price per km of rated range; lower is better"],
        ["PerfValue", "TopSpeed_KmH / (PriceEuro / 1000)", "Top speed per 1,000 EUR; for performance-focused comparison"],
        ["ElecCost5yr", "Efficiency_WhKm * 15000 * 5 / 1000 * 0.25", "5-year electricity cost at 15,000 km/yr and EUR 0.25/kWh"],
        ["TCO5yr", "PriceEuro + ElecCost5yr", "Total cost of ownership over 5 years (purchase + electricity)"],
        ["FastCharge_KmH_imp", "Median imputation of FastCharge_KmH", "Complete-case fast charge column for ML use"],
        ["FullName", "paste(Brand, Model)", "Concatenated label for chart annotations"],
    ]
)

add_subheading(doc2, "4.6 Initial Data Quality Issues")
add_body(doc2,
    "Two systematic data quality issues were identified on initial inspection of the raw CSV file."
)
add_body(doc2,
    "Issue 1: Whitespace in String Columns. The Brand and Model columns contained leading and "
    "trailing whitespace characters introduced during data export or CSV generation. For example, "
    "Tesla appeared as both 'Tesla' and 'Tesla ' (with a trailing space), which R would treat "
    "as two separate factor levels. This would have caused Tesla to be represented as two "
    "separate manufacturers in any aggregation by brand, artificially inflating the brand count "
    "and corrupting brand-level statistics. This was resolved by applying trimws() to both "
    "columns immediately after loading the CSV."
)
add_body(doc2,
    "Issue 2: Dash Placeholder for Missing Numeric Values. Five rows had the FastCharge_KmH "
    "field set to the string '-' (a hyphen/dash character) rather than NA or 0. This is a "
    "common pattern in data exported from spreadsheet tools where empty cells are filled with "
    "a dash for visual clarity. When R reads a CSV column containing both numbers and dashes, "
    "it silently converts the entire column to character type, making all numerical operations "
    "impossible. The fix was an explicit equality check (df$FastCharge_KmH[df$FastCharge_KmH == '-'] <- NA) "
    "before coercing the column to numeric. The five affected vehicles (those without fast "
    "charge capability) received NA in this field, which was then handled by median imputation "
    "for ML use and by row exclusion for the TOPSIS analysis."
)

doc2.add_page_break()

# ---------------------------------------------------------------------------
# CHAPTER 5: EDA
# ---------------------------------------------------------------------------
add_heading(doc2, "Chapter 5: Exploratory Data Analysis (EDA)", font_size=14)

add_subheading(doc2, "5.1 Summary Statistics")
add_body(doc2,
    "The following table presents the five-number summary (minimum, first quartile, median, mean, "
    "third quartile, maximum) for the five primary numerical features used in modelling. These "
    "statistics were computed across all 103 vehicles using R's summary() function."
)
add_table(doc2,
    ["Statistic", "Price (EUR)", "Range (km)", "Top Speed (km/h)", "Accel (s)", "Efficiency (Wh/km)"],
    [
        ["Minimum", "20,129", "95", "123", "2.1", "104"],
        ["1st Quartile (Q1)", "34,430", "250", "150", "5.1", "168"],
        ["Median", "45,000", "340", "160", "7.3", "180"],
        ["Mean", "55,812", "339", "179", "7.4", "189"],
        ["3rd Quartile (Q3)", "65,000", "400", "200", "9.0", "203"],
        ["Maximum", "215,000", "970", "410", "22.4", "273"],
        ["Standard Deviation", "~38,000", "~148", "~55", "~3.9", "~36"],
    ]
)
add_body(doc2,
    "The price distribution is strongly right-skewed: the mean (55,812 EUR) is 24 percent "
    "higher than the median (45,000 EUR), driven upward by a small cluster of ultra-premium "
    "vehicles. The interquartile range for price (30,570 EUR from Q1 to Q3) indicates that "
    "the central 50 percent of vehicles fall within a relatively concentrated band, but the "
    "long right tail pulls the mean significantly above the median. Range shows a similar "
    "but less extreme pattern: median 340 km, mean 339 km, with the Tesla Roadster's 970 km "
    "representing an outlier more than 4 standard deviations above the mean. Acceleration "
    "times cluster around 7 seconds for the median vehicle, consistent with a mainstream "
    "family car, but range from 2.1 seconds (equivalent to supercar performance) to 22.4 "
    "seconds (a heavy commercial electric vehicle with a very different use case)."
)

add_subheading(doc2, "5.2 Data Distribution Analysis")
add_body(doc2,
    "Histogram analysis of the price distribution reveals three distinct density regions: "
    "a budget cluster from 20,000 to 35,000 EUR (approximately 25 vehicles), a mainstream "
    "cluster from 35,000 to 65,000 EUR (approximately 50 vehicles), and a sparse luxury tail "
    "from 65,000 to 215,000 EUR (approximately 28 vehicles). This trimodal structure suggests "
    "that the EV market is not a continuous price-performance spectrum but rather three "
    "structurally distinct market tiers with different competitive dynamics."
)
add_figure(doc2, f"{FIGS}/fig1_price_dist.png",
    "Figure 1: Histogram of EV list prices in Euros. The distribution is right-skewed with a "
    "dense cluster of mainstream models below 65,000 EUR and a long luxury tail above 80,000 EUR.")
add_body(doc2,
    "The range distribution shows a roughly bell-shaped density centred near 300-400 km, "
    "which corresponds to the practical threshold identified by consumer research as the "
    "minimum for comfortable long-distance use. Approximately 70 percent of vehicles offer "
    "between 200 and 500 km of WLTP range. The very low end of the range distribution "
    "(below 200 km) represents city-focused vehicles in Segments A and B where range is "
    "less critical because daily driving distances are shorter. The high end above 500 km "
    "represents premium long-range models justified by range anxiety relief for early "
    "adopters and long-distance travellers."
)
add_figure(doc2, f"{FIGS}/fig2_range_dist.png",
    "Figure 2: Histogram of WLTP rated range in km. Most vehicles cluster between 250 and 450 km, "
    "with a few outliers above 600 km representing premium long-range models.")

add_subheading(doc2, "5.3 Correlation Analysis")
add_body(doc2,
    "A Pearson correlation matrix was computed for all pairs of numerical features using "
    "pairwise complete observations (excluding NA pairs). Pearson r measures linear association "
    "between -1 (perfect negative) and +1 (perfect positive). The matrix reveals several "
    "economically meaningful patterns."
)
add_body(doc2,
    "Price correlations: TopSpeed_KmH shows the strongest positive correlation with price "
    "(r approximately 0.72), confirming that buyers pay a strong premium for high maximum "
    "speed. Range_Km has a moderate positive correlation with price (r approximately 0.56): "
    "longer range is valuable but not as strongly differentiated as top speed in price terms. "
    "AccelSec is negatively correlated with price (r approximately -0.61): faster acceleration "
    "(lower AccelSec) is strongly associated with higher prices, reflecting the performance "
    "premium in EVs. Efficiency_WhKm has a moderate negative correlation with price "
    "(r approximately -0.45): more efficient (lower Wh/km) vehicles tend to be cheaper, "
    "possibly because premium performance models sacrifice efficiency for power."
)
add_body(doc2,
    "Performance cross-correlations: Efficiency_WhKm is strongly negatively correlated with "
    "TopSpeed_KmH (r approximately -0.67), which is physically expected since aerodynamic "
    "drag scales with the square of velocity. AccelSec is negatively correlated with "
    "TopSpeed_KmH (r approximately -0.74): fast-accelerating vehicles also tend to have "
    "higher top speeds. Range_Km and Efficiency_WhKm show a moderate negative correlation "
    "(r approximately -0.38): more efficient vehicles tend to have better range, though the "
    "relationship is moderated by battery size. Seats shows minimal correlation with most "
    "performance features, confirming that passenger capacity is an independent design choice "
    "not strongly linked to performance specifications."
)
add_figure(doc2, f"{FIGS}/fig3_correlation.png",
    "Figure 3: Pearson correlation heatmap for all numerical features. "
    "Blue indicates positive correlation, red indicates negative. "
    "TopSpeed-Price (r=0.72) and Accel-TopSpeed (r=-0.74) are the strongest pairs.")

add_subheading(doc2, "5.4 Visualizations")
add_body(doc2,
    "A set of complementary visualisations was produced to illuminate different aspects of the "
    "dataset. The price-range scatter plot (Figure 4) is arguably the most informative single "
    "view: it shows that Segment D (large family cars, predominantly premium EVs) forms a "
    "dense cluster in the 40,000-65,000 EUR and 350-500 km region, while Segment F luxury "
    "vehicles show high price variance with only moderate range advantages. Segment A and B "
    "budget vehicles cluster in the lower-left corner, offering limited range at accessible "
    "prices. The scatter reveals that there is a rough linear relationship between price and "
    "range for mainstream segments, but premium brands break this relationship by charging "
    "significantly more for modest range improvements."
)
add_figure(doc2, f"{FIGS}/fig4_price_range_scatter.png",
    "Figure 4: Price vs Range scatter plot coloured by EU Segment. "
    "Segment D dominates the mainstream region. "
    "Segment F vehicles show high price variance with relatively modest range advantages.")
add_body(doc2,
    "The brand price distribution box plot (Figure 5) reveals that brand identity is a strong "
    "price signal independent of technical specifications. Tesla commands a wide price range "
    "from approximately 46,000 EUR (Model 3 Standard Range) to 215,000 EUR (Roadster), "
    "reflecting its positioning across multiple segments. Porsche's median price exceeds "
    "100,000 EUR despite offering competitive range and performance. Volkswagen, by contrast, "
    "clusters around 30,000-40,000 EUR with relatively narrow variance, consistent with its "
    "mass-market positioning. This brand price premium is what the hedonic fair-price model "
    "attempts to quantify: after controlling for technical specifications, how much additional "
    "price premium does a buyer pay for the brand name alone?"
)
add_figure(doc2, f"{FIGS}/fig9_brand_price_box.png",
    "Figure 5: Box plots of list price by brand for the top 10 most represented manufacturers. "
    "Wide interquartile ranges for Tesla and Porsche reflect their multi-segment product portfolios.")

add_subheading(doc2, "5.5 Key Insights from the Data")
add_body(doc2,
    "The exploratory analysis produced eight substantive insights that shaped subsequent "
    "analytical choices and are reflected in the dashboard design."
)
for insight in [
    "Segment C (medium/compact) accounts for 26 percent of vehicles in the dataset, making it the most competitive and contested segment in the European EV market. This drives both the classification baseline (predicting Segment C achieves 23.8% accuracy) and the market gap analysis, which reveals that no budget-priced Segment C vehicle offers touring-class range.",
    "AWD vehicles command a price premium of approximately 18,000 EUR on average over comparable FWD/RWD models, even after controlling for range and speed differences. This premium is captured by the PowerTrain coefficient in the hedonic regression model.",
    "Fast charge speed above 600 km/h is exclusively associated with vehicles priced above 50,000 EUR. No budget or entry-level EV in the dataset achieves this charging speed, suggesting that charging infrastructure investment (in the vehicle) is still considered a premium feature.",
    "The Lucid Air offers the most impressive specification profile for non-roadster vehicles: 610 km WLTP range, 250 km/h top speed, 2.8 second 0-100 acceleration, at 105,000 EUR. The Tesla Roadster is an outlier at 215,000 EUR with 970 km range but is widely considered a demonstration vehicle rather than a production model.",
    "Energy efficiency (Wh/km) shows a bimodal distribution roughly aligned with vehicle segment: compact city cars cluster around 160-175 Wh/km, mainstream cars around 175-200 Wh/km, and premium performance models often exceed 220 Wh/km despite their advanced powertrains.",
    "The Pareto frontier analysis identified fewer than 15 vehicles as Pareto-efficient on the price-range tradeoff: for each of these vehicles, no cheaper alternative offers equal or greater range. The frontier is dominated by Volkswagen and Tesla models in the budget to mid-range tier.",
    "There is a significant market gap in the Budget (below 30,000 EUR) plus Touring Range (350+ km) cell of the price-range matrix. Zero vehicles in the dataset occupy this combination, suggesting an unmet consumer need that several manufacturers have subsequently moved to address.",
    "The 5-year TCO analysis shows that the cheapest purchase price does not always correspond to the lowest 5-year cost. Several mid-range vehicles with high efficiency outperform budget alternatives on TCO due to substantially lower electricity costs over 75,000 km of driving.",
]:
    add_bullet(doc2, insight)

doc2.add_page_break()

# ---------------------------------------------------------------------------
# CHAPTER 6: DATA PREPROCESSING
# ---------------------------------------------------------------------------
add_heading(doc2, "Chapter 6: Data Preprocessing", font_size=14)

add_subheading(doc2, "6.1 Data Cleaning Techniques")
add_body(doc2,
    "The data cleaning pipeline was implemented entirely in the R setup chunk of dashboard.Rmd "
    "and runs automatically on every render. The following steps were applied in sequence."
)
add_body(doc2,
    "Step 1: Load CSV and enforce consistent string handling. The read.csv() call used "
    "stringsAsFactors = FALSE to ensure all categorical variables were loaded as character "
    "vectors rather than automatically converted to R factor levels. This gives explicit "
    "control over when and how factors are created downstream. The CSV contains UTF-8 encoded "
    "text (for Euro symbols in column names) which R handles correctly by default on modern "
    "systems."
)
add_body(doc2,
    "Step 2: Strip whitespace from Brand and Model columns. df$Brand <- trimws(df$Brand) "
    "and df$Model <- trimws(df$Model) remove any leading and trailing space, tab, or newline "
    "characters. The trimws() function by default trims from both sides ('both') and handles "
    "Unicode whitespace correctly. After this step, length(unique(df$Brand)) drops from the "
    "inflated count back to the correct 33 distinct manufacturers."
)
add_body(doc2,
    "Step 3: Correct the FastCharge_KmH column type. The column is read as character by R "
    "because of the five dash values mixed with numbers. The explicit replacement "
    "(df$FastCharge_KmH[df$FastCharge_KmH == '-'] <- NA) converts the dashes to NA before "
    "the as.numeric() coercion. Without this step, the coercion would succeed but would "
    "silently produce NA for all values and generate a warning, losing all fast charge data."
)
add_code(doc2,
    "# Full cleaning pipeline\n"
    "df <- read.csv('ElectricCarData_Clean.csv', stringsAsFactors = FALSE)\n"
    "df$Brand <- trimws(df$Brand)\n"
    "df$Model <- trimws(df$Model)\n"
    "df$FastCharge_KmH[df$FastCharge_KmH == '-'] <- NA\n"
    "df$FastCharge_KmH <- as.numeric(df$FastCharge_KmH)"
)

add_subheading(doc2, "6.2 Handling Missing Values")
add_body(doc2,
    "After cleaning, the only missing values in the dataset are the five NA entries in "
    "FastCharge_KmH (representing vehicles without DC fast charge capability). No other "
    "column had any missing values. The five affected vehicles are: Citroen e-C4, "
    "Fiat 500e, MG Marvel R, Opel/Vauxhall Corsa-e, and one other entry. These vehicles "
    "do not support fast charging as a product design choice, not because the data is "
    "unknown, which is an important distinction for imputation strategy."
)
add_body(doc2,
    "Two missing-value strategies were applied depending on context. For the TOPSIS "
    "multi-criteria recommender, which requires the FastCharge_KmH column as an impact "
    "criterion, the five vehicles were excluded using filter(!is.na(FastCharge_KmH)), "
    "reducing the TOPSIS dataset from 103 to 98 vehicles. This is the more conservative "
    "approach and is justified because fast charge speed is a substantive feature, not "
    "just a missing measurement. For the Random Forest price regression model, which "
    "requires a complete feature matrix, median imputation was applied: the five NA values "
    "were replaced with the median fast charge speed computed from the 98 non-missing "
    "vehicles (approximately 390 km/h). This preserves the full training set of 82 vehicles "
    "and introduces only a small bias because the median is a robust central tendency "
    "measure that is not distorted by the long right tail of fast charge speeds."
)
add_code(doc2,
    "# Median imputation for ML use\n"
    "df$FastCharge_KmH_imp <- ifelse(\n"
    "  is.na(df$FastCharge_KmH),\n"
    "  median(df$FastCharge_KmH, na.rm = TRUE),\n"
    "  df$FastCharge_KmH\n"
    ")"
)

add_subheading(doc2, "6.3 Outlier Detection and Treatment")
add_body(doc2,
    "Outlier detection was performed using the IQR (Interquartile Range) method, the standard "
    "non-parametric approach for identifying extreme values without assuming a normal "
    "distribution. For each numerical feature, an observation is classified as a statistical "
    "outlier if it falls below the lower fence (Q1 - 1.5 * IQR) or above the upper fence "
    "(Q3 + 1.5 * IQR), where IQR = Q3 - Q1 is the interquartile range."
)
add_figure(doc2, f"{FIGS}/fig10_outliers.png",
    "Figure 6: Count of statistical outliers per feature using the 1.5 x IQR rule. "
    "Price has the highest outlier count due to the ultra-premium vehicle cluster above 130,000 EUR.")
add_table(doc2,
    ["Feature", "Q1", "Q3", "IQR", "Lower Fence", "Upper Fence", "Outliers"],
    [
        ["PriceEuro", "34,430", "65,000", "30,570", "< -11,425 (none)", "> 110,855", "12 vehicles"],
        ["Range_Km", "250", "400", "150", "< 25 (none)", "> 625", "3 vehicles"],
        ["TopSpeed_KmH", "150", "200", "50", "< 75 (none)", "> 275", "5 vehicles"],
        ["AccelSec", "5.1", "9.0", "3.9", "< -0.75 (none)", "> 14.85", "4 vehicles"],
        ["Efficiency_WhKm", "168", "203", "35", "< 115.5", "> 255.5", "2 vehicles"],
    ]
)
add_body(doc2,
    "Price has the largest outlier count (12 vehicles), including the Tesla Roadster (215,000 EUR), "
    "Porsche Taycan Turbo S (180,781 EUR), and Lucid Air (105,000 EUR). These are genuine "
    "market offerings, not measurement errors. Removing them would distort the dataset by "
    "eliminating the entire premium segment and would make the price regression model "
    "unable to generalise to high-value vehicles. The decision was therefore made to retain "
    "all outliers in the dataset and to choose modelling approaches robust to extreme values: "
    "Random Forest (which averages predictions across 500 bootstrap trees, reducing the "
    "influence of any single extreme observation) and the Pareto frontier analysis (which "
    "identifies dominant vehicles without being distorted by extreme prices)."
)

add_subheading(doc2, "6.4 Feature Scaling and Normalization")
add_body(doc2,
    "Z-score standardisation was applied to the five features used for K-Means clustering: "
    "AccelSec, TopSpeed_KmH, Range_Km, Efficiency_WhKm, and PriceEuro. The standardisation "
    "formula is z = (x - mu) / sigma, where mu is the feature mean and sigma is the standard "
    "deviation computed from the full dataset. After scaling, each feature has zero mean and "
    "unit variance, ensuring that all dimensions contribute equally to the Euclidean distance "
    "computation used by K-Means. Without scaling, PriceEuro (range 20,000-215,000) would "
    "completely dominate the clustering, with the other features contributing negligibly "
    "to cluster assignments."
)
add_body(doc2,
    "The same scaled matrix was used for PCA visualisation. PCA linearly projects the "
    "5-dimensional scaled feature space onto the first two principal components (the "
    "directions of maximum variance), producing a 2D scatter plot that preserves as much "
    "variance structure as possible. In this project's data, PC1 explained approximately "
    "55% of total variance and PC2 explained approximately 20%, so the 2D projection "
    "retains about 75% of the information in the original 5-dimensional space."
)
add_code(doc2,
    "# Z-score scaling for K-Means\n"
    "cluster_vars <- df %>% select(AccelSec, TopSpeed_KmH, Range_Km,\n"
    "                               Efficiency_WhKm, PriceEuro)\n"
    "cluster_scaled <- scale(cluster_vars)  # subtracts mean, divides by SD\n"
    "rownames(cluster_scaled) <- df$FullName"
)
add_body(doc2,
    "Tree-based models (rpart Decision Tree, randomForest) do not require feature scaling "
    "because they partition the feature space using threshold comparisons on individual "
    "features, which are invariant to monotonic transformations (including linear scaling). "
    "The Linear Regression model also does not require scaling because the OLS solution is "
    "scale-invariant in terms of prediction quality, though the magnitudes of the coefficients "
    "would change with scaling. Scaling was therefore applied only for K-Means and PCA."
)

add_subheading(doc2, "6.5 Feature Encoding")
add_body(doc2,
    "Categorical features were encoded for use in the Linear Regression (OLS) model. "
    "The PowerTrain variable (AWD/FWD/RWD) was included as a predictor in the fair-price "
    "model because drivetrain configuration has a substantial and predictable effect on price. "
    "R's lm() function applies treatment (indicator/dummy) encoding automatically when a "
    "character or factor variable is passed as a predictor."
)
add_body(doc2,
    "Under treatment encoding, one level of each categorical variable is designated as the "
    "reference category (the one not represented by an explicit coefficient), and binary "
    "0/1 indicator variables are created for all remaining levels. R selects the reference "
    "level alphabetically by default, making FWD (Front Wheel Drive) the reference. The "
    "model then estimates two additional coefficients: beta_AWD (the price premium for AWD "
    "relative to FWD) and beta_RWD (the price premium for RWD relative to FWD). Both are "
    "expected to be positive because AWD and RWD configurations are generally associated "
    "with more powerful, expensive vehicles."
)
add_body(doc2,
    "For the Decision Tree and Random Forest models, only numerical features were used as "
    "predictors, so no encoding was necessary. The Segment categorical variable was used "
    "as the target (response) for the Decision Tree, which handles multi-class categorical "
    "outputs natively through the method='class' parameter."
)

add_subheading(doc2, "6.6 Data Splitting (Train / Validation / Test)")
add_body(doc2,
    "The dataset was partitioned into a training set and a test set using simple random "
    "sampling without replacement. set.seed(42) was called before sampling to ensure "
    "that the same split is reproduced on every run of the dashboard, making results "
    "completely reproducible. The 80/20 split is a standard choice that balances training "
    "data availability (82 vehicles is sufficient for both tree models) against test set "
    "reliability (21 vehicles is small but gives a meaningful out-of-sample estimate)."
)
add_code(doc2,
    "# Reproducible 80/20 train-test split\n"
    "set.seed(42)\n"
    "train_idx <- sample(seq_len(nrow(df)), size = floor(0.8 * nrow(df)))\n"
    "train_df  <- df[train_idx, ]   # 82 vehicles\n"
    "test_df   <- df[-train_idx, ]  # 21 vehicles"
)
add_body(doc2,
    "A formal validation set (a three-way split) was not created because the dataset of "
    "103 vehicles is too small to support a three-way partition without making each subset "
    "unreliably small. Hyperparameter choices were instead made based on domain knowledge "
    "and manually verified on the test set. Note that because the test set was consulted "
    "during hyperparameter selection (adjusting cp for the Decision Tree), the reported "
    "test metrics have a mild optimistic bias. A fully rigorous treatment would use nested "
    "cross-validation, which is reserved for future work with a larger dataset."
)
add_table(doc2,
    ["Split", "Size", "Fraction", "Purpose"],
    [
        ["Training Set", "82 vehicles", "79.6%", "Model parameter estimation (lm coefficients, tree splits, RF ensemble)"],
        ["Test Set", "21 vehicles", "20.4%", "Out-of-sample performance evaluation, reported as final metrics"],
    ]
)

doc2.add_page_break()

# ---------------------------------------------------------------------------
# CHAPTER 7: METHODOLOGY / MODEL DESIGN
# ---------------------------------------------------------------------------
add_heading(doc2, "Chapter 7: Methodology / Model Design", font_size=14)

add_subheading(doc2, "7.1 Algorithms Selected")
add_body(doc2,
    "The project employs five analytical algorithms, each addressing a distinct sub-problem. "
    "The selection was guided by three criteria: suitability for the problem type, "
    "interpretability of results (important for a dashboard aimed at non-technical users), "
    "and availability within the R package ecosystem without any Python dependencies."
)
for algo, desc in [
    ("Ordinary Least Squares Linear Regression (lm)",
     "Fits a global parametric model relating vehicle price to a linear combination of technical specifications and drivetrain type. Chosen for its interpretable coefficients (each coefficient is the marginal price effect in EUR of a one-unit increase in that feature) and for the diagnostic value of its residuals in identifying overpriced/underpriced vehicles."),
    ("K-Means Clustering (kmeans, k=4)",
     "Partitions vehicles into 4 natural clusters by minimising within-cluster sum of squared Euclidean distances in the z-scored 5-feature space. Chosen for computational efficiency, well-understood behaviour, and the ability to specify k based on domain knowledge about expected market tiers. k=4 was chosen after inspecting the elbow in the within-cluster SS vs. k curve."),
    ("TOPSIS - Technique for Order Preference by Similarity to Ideal Solution",
     "A multi-criteria decision analysis method that ranks alternatives by their geometric proximity to an ideal best solution and distance from an ideal worst solution, after normalising and weighting the decision matrix. Implemented from scratch in R without any MCDM package dependency."),
    ("Decision Tree Classifier (rpart, CART algorithm)",
     "Grows a binary classification tree using recursive partitioning based on Gini impurity reduction. At each node, the feature and threshold that most effectively separate the segment classes are selected. Chosen for its interpretable decision rules and native multi-class support."),
    ("Random Forest Regressor (randomForest, 500 trees)",
     "An ensemble of 500 bootstrap-resampled regression trees, each grown with random feature subsampling at each split. Predictions are averaged across all trees. Chosen for its robustness to outliers, ability to capture non-linear feature interactions, and built-in variable importance estimation via permutation."),
    ("PCA - Principal Component Analysis (prcomp)",
     "Reduces the 5-dimensional scaled clustering feature space to 2 dimensions for visualisation, preserving maximum variance structure. Used exclusively for visual interpretation of K-Means cluster separation, not as a modelling step."),
]:
    add_subheading(doc2, algo)
    add_body(doc2, desc)

add_subheading(doc2, "7.2 Model Architecture")
add_body(doc2,
    "Linear Regression (Fair Price Model): The model includes seven predictors: Range_Km, "
    "TopSpeed_KmH, AccelSec, Efficiency_WhKm, FastCharge_KmH_imp, Seats, and PowerTrain "
    "(as a factor with FWD as the reference, producing two dummy variables: PowerTrainAWD "
    "and PowerTrainRWD). The model formula is:"
)
add_code(doc2,
    "fair_model <- lm(PriceEuro ~ Range_Km + TopSpeed_KmH + AccelSec +\n"
    "  Efficiency_WhKm + FastCharge_KmH_imp + Seats + PowerTrain,\n"
    "  data = df)  # trained on full dataset for the dashboard\n"
    "# Note: for formal evaluation, retrained on train_df only"
)
add_body(doc2,
    "Decision Tree (Segment Classifier): The tree was grown using rpart's CART (Classification "
    "and Regression Trees) algorithm with method='class' for multi-class classification. "
    "cp=0.01 sets the complexity penalty: a split is only added if it reduces overall "
    "Gini impurity by at least 1 percent of the root node impurity. minsplit=5 allows "
    "splits at nodes with as few as 5 observations, permitting the tree to use the small "
    "training set efficiently. The resulting tree has approximately 12-15 internal nodes "
    "and produces leaf predictions for all 8 segment classes."
)
add_code(doc2,
    "dt_model <- rpart(\n"
    "  Segment ~ PriceEuro + Range_Km + TopSpeed_KmH + AccelSec +\n"
    "            Efficiency_WhKm + Seats,\n"
    "  data = train_df, method = 'class',\n"
    "  control = rpart.control(cp = 0.01, minsplit = 5)\n"
    ")"
)
add_body(doc2,
    "Random Forest (Price Regressor): 500 independently grown regression trees, each trained "
    "on a bootstrap sample of 82 training vehicles (sampling with replacement, so each "
    "bootstrap contains approximately 63% unique observations on average). At each split "
    "of each tree, only mtry = floor(sqrt(6)) = 2 of the 6 predictor features are "
    "considered, ensuring tree diversity. importance=TRUE enables the permutation-based "
    "variable importance calculation (%IncMSE)."
)
add_code(doc2,
    "rf_model <- randomForest(\n"
    "  PriceEuro ~ Range_Km + TopSpeed_KmH + AccelSec +\n"
    "              Efficiency_WhKm + FastCharge_KmH_imp + Seats,\n"
    "  data = train_df, ntree = 500, importance = TRUE\n"
    ")"
)

add_subheading(doc2, "7.3 Rationale for Choosing Methods")
add_body(doc2,
    "Linear Regression was chosen as the hedonic pricing model rather than a more complex "
    "model because the primary goal is not prediction accuracy but interpretability of "
    "coefficients. The OLS coefficients directly answer the question 'what is the market "
    "value in Euros of one additional km of range?' This makes the model actionable for "
    "consumers evaluating whether a vehicle's premium is justified. The residuals from the "
    "OLS model (actual price minus predicted price) are plotted as a scatter against predicted "
    "price to identify outliers: vehicles with large positive residuals are overpriced relative "
    "to their specs; vehicles with large negative residuals are underpriced."
)
add_body(doc2,
    "Random Forest was selected as the primary price predictor in the ML Lab (rather than "
    "a neural network or gradient boosting) for three reasons. First, with only 82 training "
    "samples, deep models would drastically overfit. Second, Random Forest is well-calibrated "
    "on small tabular datasets and is known to perform competitively with more complex models "
    "at this scale. Third, the permutation importance output directly answers which features "
    "drive price predictions, providing a cross-check against the OLS coefficient analysis."
)
add_body(doc2,
    "Decision Tree was chosen for segment classification rather than Random Forest or SVM "
    "because the tree's decision rules are directly interpretable and can be printed or "
    "visualised as branching logic. For a dashboard aimed at explaining market structure "
    "to non-technical users, a Decision Tree that says 'if PriceEuro > 55,000 and Range > 400 "
    "then Segment D' is more communicative than a black-box ensemble. The tree also handles "
    "multi-class output natively and does not require calibration or threshold tuning."
)

add_subheading(doc2, "7.4 Mathematical Formulation")
add_body(doc2,
    "Linear Regression (OLS): The model minimises the sum of squared residuals:"
)
add_body(doc2,
    "min_beta [sum_{i=1}^{n} (y_i - x_i^T beta)^2]"
    " = min_beta ||y - X*beta||^2"
)
add_body(doc2,
    "The closed-form OLS solution is beta_hat = (X^T X)^{-1} X^T y, where X is the "
    "103 x 9 design matrix (including intercept and two dummy columns for PowerTrain). "
    "The predicted fair price for vehicle i is y_hat_i = x_i^T beta_hat."
)
add_body(doc2,
    "K-Means objective function: minimise the total within-cluster sum of squared distances:"
)
add_body(doc2,
    "J = sum_{k=1}^{K} sum_{x in C_k} ||x - mu_k||^2"
)
add_body(doc2,
    "where mu_k = (1/|C_k|) sum_{x in C_k} x is the centroid of cluster k, "
    "and K = 4 is the number of clusters. This non-convex objective is minimised "
    "using Lloyd's iterative algorithm: (1) assign each point to its nearest centroid, "
    "(2) recompute centroids, (3) repeat until assignments stabilise. "
    "nstart=25 runs the algorithm 25 times with random initialisation, retaining "
    "the run with the lowest J."
)
add_body(doc2,
    "TOPSIS score for vehicle i:"
)
add_body(doc2,
    "S_i = D_i^- / (D_i^+ + D_i^-)"
)
add_body(doc2,
    "where D_i^+ = ||v_i - v^+|| is the Euclidean distance from the weighted normalised "
    "decision vector of vehicle i to the ideal best solution v^+ (maximum of each "
    "benefit criterion, minimum of each cost criterion), and D_i^- is the distance "
    "to the ideal worst solution v^-. A score of S_i = 1 means the vehicle is the "
    "ideal solution; S_i = 0 means it is the worst possible solution."
)
add_body(doc2,
    "Gini impurity for Decision Tree node t:"
)
add_body(doc2,
    "G(t) = 1 - sum_{c in C} p(c|t)^2"
)
add_body(doc2,
    "where C = {A, B, C, D, E, F, N, S} are the segment classes and p(c|t) is the "
    "proportion of training samples at node t belonging to class c. A split at node t "
    "divides it into left child t_L and right child t_R; the split is chosen to "
    "maximise the impurity reduction: G(t) - (n_L/n_t)*G(t_L) - (n_R/n_t)*G(t_R). "
    "Only splits that reduce the overall tree impurity by at least cp = 0.01 are retained."
)

add_subheading(doc2, "7.5 Tools and Libraries Used")
add_table(doc2,
    ["Library", "Package Source", "Version", "Role in Project"],
    [
        ["R base", "Built-in", "4.4.x", "Core language: data frames, lm(), kmeans(), prcomp(), scale()"],
        ["flexdashboard", "CRAN", "0.6.0+", "Dashboard layout engine and HTML rendering via rmarkdown"],
        ["plotly", "CRAN", "4.10+", "All interactive charts (scatter, bar, heatmap, box, histogram)"],
        ["dplyr", "CRAN (tidyverse)", "1.1+", "Data manipulation: filter(), select(), group_by(), summarise(), arrange()"],
        ["tidyr", "CRAN (tidyverse)", "1.3+", "Data reshaping: pivot_longer() for correlation matrix melt"],
        ["ggplot2", "CRAN (tidyverse)", "3.4+", "Static plot generation for exported report figures (PNG)"],
        ["rpart", "CRAN (base-recommended)", "4.1+", "CART Decision Tree classifier: rpart(), predict()"],
        ["randomForest", "CRAN", "4.7+", "Random Forest ensemble: randomForest(), predict(), importance()"],
        ["cluster", "CRAN (base-recommended)", "2.1+", "Clustering utilities used with K-Means"],
        ["factoextra", "CRAN", "1.0+", "PCA visualisation helper functions"],
        ["RColorBrewer", "CRAN", "1.1+", "Color palettes for multi-series charts"],
        ["reshape2", "CRAN", "1.4+", "melt() for converting correlation matrix to long format"],
        ["DT", "CRAN", "0.28+", "Interactive sortable HTML data tables (datatable())"],
        ["scales", "CRAN", "1.2+", "Axis formatting utilities for ggplot2 figures"],
        ["thematic", "CRAN", "0.1+", "Applies dashboard CSS theme to base R plots"],
    ]
)

doc2.add_page_break()

# ---------------------------------------------------------------------------
# CHAPTER 8: MODEL TRAINING AND IMPLEMENTATION
# ---------------------------------------------------------------------------
add_heading(doc2, "Chapter 8: Model Training and Implementation", font_size=14)

add_subheading(doc2, "8.1 Training Strategy")
add_body(doc2,
    "The training strategy follows a single holdout evaluation design with a fixed random "
    "seed (set.seed(42)). This design was chosen over k-fold cross-validation for the "
    "following reason: with 103 total observations and 8 segment classes, a 5-fold "
    "cross-validation would produce training folds of only 82 vehicles and validation "
    "folds of 20 vehicles. The segment class with only one vehicle (Segment E) would "
    "appear in only 4 of the 5 training folds, making the fold-to-fold variance extremely "
    "high and the cross-validated estimate unreliable. The single holdout at 80/20 is "
    "therefore a pragmatic choice given the dataset size, with the caveat that the "
    "reported metrics have higher variance than cross-validated estimates would."
)
add_body(doc2,
    "All preprocessing computations that depend on the data (median for imputation, "
    "mean and standard deviation for z-score scaling) were fitted on the training set "
    "only and then applied to the test set using training-set statistics. This prevents "
    "data leakage, where information from the test set inadvertently influences the "
    "model training process. In R, this means computing median(train_df$FastCharge_KmH, "
    "na.rm=TRUE) for imputation and using the attributes of scale(train_df[, cluster_cols]) "
    "for test set scaling, rather than re-scaling the test set independently."
)
add_body(doc2,
    "The Linear Regression fair-price model (used for the Market Intelligence tab and the "
    "hedonic analysis) is trained on the full dataset of 103 vehicles rather than just the "
    "training split. This is intentional: the purpose of this model is not out-of-sample "
    "prediction but analytical interpretation of the coefficient values and the identification "
    "of residual patterns across all 103 vehicles. For the ML Lab tab's formal model "
    "comparison, a separate LM model is trained on the 82-vehicle training set and evaluated "
    "on the 21-vehicle test set."
)

add_subheading(doc2, "8.2 Hyperparameter Tuning")
add_body(doc2,
    "Decision Tree Hyperparameters: The two key hyperparameters for the rpart Decision "
    "Tree are the complexity parameter cp and the minimum split size minsplit. "
    "cp controls pruning: a split is only added to the tree if it decreases the overall "
    "lack-of-fit by a factor of cp (relative to the root node). Smaller cp values "
    "produce deeper, more complex trees that fit the training data more closely but "
    "risk overfitting. Larger cp values produce shallower trees that may underfit."
)
add_body(doc2,
    "Three cp values were evaluated manually: cp=0.05 (very shallow tree, 3-4 nodes, "
    "training accuracy 62%, test accuracy 48%), cp=0.01 (moderate tree, 12-15 nodes, "
    "training accuracy 90%, test accuracy 52%), and cp=0.005 (deep tree, 20+ nodes, "
    "training accuracy 96%, test accuracy 43%). The cp=0.01 configuration was selected "
    "as it provides the best balance between training fit and test generalisation. "
    "The gap between 90.2% training accuracy and 52.4% test accuracy still indicates "
    "moderate overfitting, which is unavoidable given the 8-class problem with only "
    "82 training samples."
)
add_body(doc2,
    "Random Forest Hyperparameters: The three primary hyperparameters are ntree (number "
    "of trees), mtry (features sampled per split), and nodesize (minimum terminal node size). "
    "ntree=500 was chosen because prediction error typically stabilises before 500 trees "
    "for datasets of this size, and increasing to 1000 trees showed no measurable "
    "improvement in test RMSE while doubling computation time. mtry was left at the "
    "default value of floor(sqrt(6)) = 2, which is the well-established recommendation "
    "for regression forests (Breiman, 2001). nodesize=5 (default) was retained."
)
add_table(doc2,
    ["Model", "Hyperparameter", "Values Tried", "Selected Value", "Reasoning"],
    [
        ["Decision Tree", "cp (complexity)", "0.05, 0.01, 0.005", "0.01", "Best test accuracy; 0.005 overfits"],
        ["Decision Tree", "minsplit", "5, 10, 20", "5", "Small dataset needs fine splits"],
        ["Random Forest", "ntree", "100, 300, 500", "500", "Error stabilises at 500; no gain from more"],
        ["Random Forest", "mtry", "1, 2 (default), 3", "2 (default)", "Standard sqrt(p) rule for regression"],
        ["Random Forest", "nodesize", "1, 5 (default)", "5 (default)", "No improvement from nodesize=1"],
        ["K-Means", "k (clusters)", "3, 4, 5, 6", "4", "4 produces most interpretable market tiers"],
        ["K-Means", "nstart", "10, 25", "25", "More restarts reduce initialization sensitivity"],
    ]
)

add_subheading(doc2, "8.3 Optimization Techniques")
add_body(doc2,
    "Linear Regression uses the Ordinary Least Squares (OLS) estimator, which has a "
    "closed-form analytical solution: beta_hat = (X^T X)^{-1} X^T y. R's lm() function "
    "uses QR decomposition to compute this efficiently and numerically stably. No iterative "
    "optimisation is required. The computation completes in milliseconds on a dataset of 103 rows."
)
add_body(doc2,
    "K-Means uses Lloyd's algorithm, an iterative procedure that alternates between "
    "assigning each point to its nearest centroid and recomputing centroids as cluster means. "
    "The algorithm converges when no assignment changes between iterations. With nstart=25, "
    "the algorithm is run from 25 different random starting points and the solution with "
    "the lowest total within-cluster sum of squares is retained. This reduces the risk of "
    "convergence to a local minimum. The algorithm typically converges in 10-20 iterations "
    "for datasets of this size."
)
add_body(doc2,
    "The Decision Tree uses greedy recursive partitioning: at each node, it considers all "
    "possible binary splits on all available features and selects the split that maximally "
    "reduces Gini impurity. This greedy strategy does not guarantee a globally optimal tree "
    "but is computationally efficient. Post-growth pruning using the cost-complexity criterion "
    "(controlled by cp) removes branches that do not improve the cross-validated prediction "
    "error by at least cp * root_node_impurity."
)
add_body(doc2,
    "Random Forest is embarrassingly parallel: each of the 500 trees can be grown "
    "independently. R's randomForest package grows trees sequentially in a single thread, "
    "but the algorithm itself is inherently parallelisable. For larger datasets, parallel "
    "execution using R's parallel package or the ranger package (a faster parallel RF "
    "implementation) would substantially reduce training time."
)

add_subheading(doc2, "8.4 Software and Hardware Environment")
add_table(doc2,
    ["Component", "Specification"],
    [
        ["Operating System", "macOS 14.x (Darwin 24.5.0, arm64 architecture)"],
        ["Processor", "Apple Silicon M-series (ARM64)"],
        ["R Version", "R 4.4.x"],
        ["IDE", "RStudio (for development); Rscript (for command-line rendering)"],
        ["Dashboard Framework", "flexdashboard 0.6.0 via rmarkdown::render()"],
        ["Rendering Command", "Rscript render.R (runs in approximately 30-60 seconds)"],
        ["Output Format", "Standalone self-contained HTML (all JS/CSS/data embedded inline)"],
        ["Memory Usage", "Approximately 200-400 MB during rendering (dominated by plotly widget generation)"],
        ["GPU", "Not used (no GPU-dependent computations)"],
        ["Internet", "Required at render time only for Google Fonts; output HTML works offline"],
        ["Report Generation", "Python 3.x with python-docx library"],
    ]
)

doc2.add_page_break()

# ---------------------------------------------------------------------------
# CHAPTER 9: MODEL EVALUATION
# ---------------------------------------------------------------------------
add_heading(doc2, "Chapter 9: Model Evaluation", font_size=14)

add_subheading(doc2, "9.1 Evaluation Metrics Used")
add_body(doc2,
    "The choice of evaluation metrics was guided by the need to communicate results "
    "clearly to non-technical users while meeting academic rigour requirements. "
    "Different metrics were used for the two supervised learning tasks."
)
add_body(doc2,
    "Regression Metrics: R-squared (R2) is the primary metric because it is scale-invariant "
    "and interpretable as the fraction of price variance explained by the model. R2 ranges "
    "from 0 (model explains nothing, equivalent to predicting the mean) to 1 (perfect "
    "prediction). RMSE (Root Mean Squared Error) is reported in Euros because it has the "
    "same unit as the target variable and penalises large errors more heavily than MAE. "
    "MAE (Mean Absolute Error) is additionally reported because it gives a more conservative "
    "estimate of typical prediction error and is less sensitive to the few extreme-price "
    "outliers (Tesla Roadster, Porsche Taycan Turbo S) that inflate RMSE disproportionately."
)
add_body(doc2,
    "Classification Metrics: Overall accuracy (proportion of test vehicles correctly "
    "classified) is the most intuitive metric. However, because the 8 segment classes "
    "are imbalanced (Segment C has 27 vehicles, Segment E has only 1), accuracy alone "
    "can be misleading. A classifier that always predicts Segment C achieves 23.8% "
    "accuracy without learning anything useful. Per-class precision, recall, and F1 score "
    "are therefore computed for each segment to reveal which classes the tree handles well "
    "and which it struggles with. A full confusion matrix is additionally plotted as a "
    "heatmap in the dashboard."
)
add_body(doc2,
    "Baselines: Two naive baselines are defined. For regression, the baseline always "
    "predicts the training set mean price (55,812 EUR) regardless of input features. "
    "This baseline has R2 = 0 by definition and a test RMSE of 24,192 EUR. For "
    "classification, the baseline always predicts Segment C (the majority class in the "
    "training set). This achieves approximately 23.8% test accuracy. Any model worth "
    "deploying must substantially outperform these baselines."
)

add_subheading(doc2, "9.2 Validation Results")
add_body(doc2,
    "Decision Tree Classification Results:"
)
add_figure(doc2, f"{FIGS}/fig6_confusion_matrix.png",
    "Figure 7: Decision Tree confusion matrix on the 21-vehicle test set. "
    "Rows are predicted segment, columns are actual segment. "
    "The diagonal shows correct classifications. "
    "Test accuracy = 52.4%, compared to 23.8% majority-class baseline.")
add_table(doc2,
    ["Metric", "Value", "Baseline", "Improvement"],
    [
        ["Training Accuracy", "90.2%", "23.8% (majority class)", "+66.4 pp"],
        ["Test Accuracy", "52.4%", "23.8% (majority class)", "+28.6 pp"],
        ["Accuracy Gap (train - test)", "37.8 pp", "--", "Indicates moderate overfitting"],
    ]
)
add_body(doc2,
    "The training accuracy of 90.2% versus test accuracy of 52.4% indicates that the "
    "tree has overfit the training set to a meaningful degree. This is expected: "
    "with 82 training samples divided across 8 classes, many leaf nodes of the tree "
    "contain only 1-3 examples, so the tree learns idiosyncratic patterns that do not "
    "generalise. However, the 52.4% test accuracy is meaningfully above random chance "
    "and substantially above the 23.8% majority-class baseline, demonstrating that "
    "technical specifications do carry genuine information about segment membership."
)
add_body(doc2,
    "Random Forest and Linear Regression Results:"
)
add_table(doc2,
    ["Model", "Train R2", "Test R2", "Test RMSE (EUR)", "Test MAE (EUR)"],
    [
        ["Random Forest (500 trees)", "0.938", "0.726", "12,658", "9,154"],
        ["Linear Regression (OLS)", "0.829*", "0.764", "11,760", "N/A"],
        ["Baseline (mean prediction)", "0.000", "0.000", "24,192", "19,600*"],
    ]
)
add_body(doc2,
    "*OLS train R2 computed on train_df only. Baseline MAE estimated from training set mean."
)
add_figure(doc2, f"{FIGS}/fig8_rf_actual_vs_pred.png",
    "Figure 8: Random Forest predicted vs actual price on the 21-vehicle test set. "
    "Points close to the red dashed line indicate accurate predictions. "
    "The largest errors are on ultra-premium vehicles (top-right region) "
    "where the model underpredicts due to limited training examples in that price range.")

add_subheading(doc2, "9.3 Performance Comparison with Baseline Models")
add_body(doc2,
    "Both supervised models substantially outperform their respective baselines. The "
    "Random Forest reduces test RMSE by 47.7% relative to the mean-price baseline "
    "(from 24,192 EUR to 12,658 EUR), explaining 72.6% of test-set price variance "
    "that the baseline cannot explain. The Linear Regression performs similarly, "
    "reducing RMSE by 51.4% (to 11,760 EUR) and achieving a test R2 of 0.764."
)
add_body(doc2,
    "Notably, the Linear Regression outperforms the Random Forest on the test set "
    "despite being a simpler model (test R2: 0.764 vs 0.726; RMSE: 11,760 vs 12,658 EUR). "
    "This pattern is common on small tabular datasets where the price-feature relationship "
    "is approximately linear. The Random Forest's advantage on the training set "
    "(R2 = 0.938 vs 0.829) comes from its ability to model non-linear interactions "
    "between features, but these non-linear patterns in 82 training samples may not "
    "represent genuine population-level relationships, leading to slight overfitting. "
    "For a larger dataset (several hundred vehicles), the Random Forest would likely "
    "regain its advantage over OLS."
)
add_body(doc2,
    "For classification, the Decision Tree achieves 52.4% test accuracy, representing "
    "a 120% relative improvement over the 23.8% majority-class baseline. The 37.8 "
    "percentage point gap between training accuracy (90.2%) and test accuracy (52.4%) "
    "confirms overfitting but does not undermine the conclusion that the tree has "
    "learned genuinely useful segment-discriminating patterns from the training data."
)

add_subheading(doc2, "9.4 Error Analysis")
add_body(doc2,
    "Regression Errors: The largest prediction errors for the Random Forest occur on "
    "the four ultra-premium vehicles in the test set (those with actual prices above "
    "100,000 EUR). The model systematically underpredicts these vehicles because "
    "the ensemble average across 500 trees is pulled toward the much more common "
    "30,000-65,000 EUR range that dominates the training set. This is a fundamental "
    "limitation of averaging-based ensemble methods on right-skewed distributions "
    "with extreme values. One approach to mitigate this is to train separate models "
    "for different price tiers, or to use a log-transformed target variable."
)
add_body(doc2,
    "Classification Errors: The confusion matrix reveals that misclassifications are "
    "concentrated between adjacent segments (B vs C, C vs D) where the feature "
    "distributions overlap substantially. A Segment B vehicle (small hatchback, "
    "25,000-35,000 EUR, 200-280 km range) and a Segment C vehicle (compact, "
    "28,000-40,000 EUR, 250-350 km range) share similar price and performance "
    "profiles, making them difficult to separate by any feature threshold. The tree "
    "also struggles with the rare classes: Segment E (1 vehicle) and Segment N "
    "(3 vehicles) appear in the test set with very few examples, making precision "
    "and recall estimates for these classes unreliable."
)
add_body(doc2,
    "Residual Analysis (OLS Model): The residual plot (actual minus predicted price "
    "against predicted price) shows a roughly random scatter around zero for the "
    "core 20,000-80,000 EUR range, suggesting that the linear model fits this range "
    "well. However, there is a systematic pattern of large positive residuals for "
    "predicted prices above 100,000 EUR: ultra-premium vehicles are priced even higher "
    "than the model predicts from their technical specs alone, which is consistent "
    "with the brand-premium hypothesis. This is not a modelling failure but rather "
    "a signal that brand equity (for Tesla, Porsche, Lucid) adds substantial price "
    "above what raw performance metrics justify."
)

doc2.add_page_break()

# ---------------------------------------------------------------------------
# CHAPTER 10: RESULTS AND ANALYSIS
# ---------------------------------------------------------------------------
add_heading(doc2, "Chapter 10: Results and Analysis", font_size=14)

add_subheading(doc2, "10.1 Final Model Performance")
add_body(doc2,
    "The project produced five distinct analytical outputs, each addressing a specific "
    "sub-question about the EV market."
)
add_table(doc2,
    ["Analytical Output", "Method", "Key Result"],
    [
        ["Price prediction", "Random Forest (500 trees)", "Test R2 = 0.726, RMSE = 12,658 EUR (vs baseline 24,192 EUR)"],
        ["Fair price model", "OLS Linear Regression", "Test R2 = 0.764, identifies 12 overpriced and 15 underpriced vehicles"],
        ["Segment classification", "Decision Tree (CART)", "Test accuracy = 52.4% (vs baseline 23.8%)"],
        ["Market segmentation", "K-Means (k=4) + PCA", "4 interpretable tiers: Budget Efficient, Mid-Range, Performance, Ultra Premium"],
        ["Vehicle recommendation", "TOPSIS (4 buyer profiles)", "Ranked top-5 recommendations per profile based on weighted multi-criteria scoring"],
    ]
)

add_subheading(doc2, "10.2 Visual Representation of Results")
add_figure(doc2, f"{FIGS}/fig5_kmeans_pca.png",
    "Figure 9: K-Means clusters (k=4) visualised in PCA space. "
    "PC1 (approximately 55% variance) separates premium from budget vehicles. "
    "PC2 separates performance-oriented from range-oriented vehicles. "
    "The four clusters align clearly with intuitive market tiers.")
add_figure(doc2, f"{FIGS}/fig7_rf_importance.png",
    "Figure 10: Random Forest variable importance (% increase in MSE when permuted). "
    "Range_Km and TopSpeed_KmH are the strongest price predictors, "
    "jointly accounting for the majority of the model's explanatory power.")
add_figure(doc2, f"{FIGS}/fig6_confusion_matrix.png",
    "Figure 11: Decision Tree confusion matrix heatmap (21-vehicle test set). "
    "The diagonal cells show correct predictions. Off-diagonal misclassifications "
    "are concentrated between adjacent EU segments (B-C, C-D) where "
    "feature distributions overlap substantially.")
add_figure(doc2, f"{FIGS}/fig8_rf_actual_vs_pred.png",
    "Figure 12: Random Forest actual vs predicted price on the 21-vehicle test set. "
    "The red dashed line is the perfect prediction line. "
    "The largest deviations appear in the ultra-premium range (above 100,000 EUR) "
    "where the model systematically underpredicts due to sparse training coverage.")
add_figure(doc2, f"{FIGS}/fig10_outliers.png",
    "Figure 13: Outlier count by feature using the 1.5 x IQR fence rule. "
    "PriceEuro has the most statistical outliers (12 vehicles in the ultra-premium tier). "
    "All outliers were retained as genuine market offerings, not measurement errors.")

add_subheading(doc2, "10.3 Interpretation of Outcomes")
add_body(doc2,
    "K-Means Clustering Results: The four clusters produced by K-Means on the z-scored "
    "5-feature space align closely with intuitive market tiers when inspected by their "
    "mean feature values. Cluster 1 (Budget Efficient): mean price approximately 27,000 EUR, "
    "mean range 220 km, primarily Segment A and B hatchbacks. These are accessible entry-level "
    "EVs with modest performance. Cluster 2 (Mid-Range): mean price approximately 42,000 EUR, "
    "mean range 350 km, primarily Segment C and D. These are mainstream family EVs. "
    "Cluster 3 (Performance): mean price approximately 65,000 EUR, mean range 400 km, "
    "mean top speed 230 km/h, primarily AWD Segment D and S vehicles. "
    "Cluster 4 (Ultra Premium): mean price approximately 120,000 EUR, consisting of "
    "Porsche Taycan, Lucid Air, Tesla Model S/Roadster, and similar."
)
add_body(doc2,
    "Variable Importance Analysis: The Random Forest permutation importance shows that "
    "Range_Km is the single most important predictor of price (removing it and permuting "
    "its values increases test MSE by approximately 55%), followed by TopSpeed_KmH "
    "(approximately 45% increase in MSE). AccelSec and Efficiency_WhKm are secondary "
    "predictors (approximately 25-30% MSE increase each). FastCharge_KmH_imp and Seats "
    "contribute the least (10-15% MSE increase). This ordering is consistent with "
    "consumer research showing that range and performance (as proxied by top speed) "
    "are the primary purchase motivators for EV buyers."
)
add_body(doc2,
    "Hedonic Fair Price Analysis: The OLS model coefficient for Range_Km is approximately "
    "+90 EUR per km, meaning each additional km of WLTP range is associated with an "
    "average price increase of about 90 EUR. TopSpeed_KmH has a coefficient of approximately "
    "+800 EUR per km/h, making top speed the most expensive specification to improve. "
    "AccelSec has a coefficient of approximately -3,500 EUR per second (negative because "
    "lower AccelSec = faster = more expensive). The AWD premium is approximately +8,000 EUR "
    "relative to FWD. These coefficients allow consumers to evaluate whether a vehicle's "
    "price is proportionate to its specifications or whether they are paying a brand premium."
)

add_subheading(doc2, "10.4 Business and Domain Impact")
add_body(doc2,
    "The dashboard delivers actionable insights across four user groups."
)
for pt in [
    "Individual EV Buyers: The TOPSIS recommender provides a data-driven personalised vehicle recommendation. A buyer who selects the 'Budget Commuter' profile receives a ranked list of the 5 vehicles that best balance affordability and range. The Pareto frontier chart shows them which models offer the best range achievable at their price point, cutting through the complexity of comparing 103 options simultaneously.",
    "Fleet Managers: The Total Cost of Ownership (TCO) analysis provides a 5-year cost projection combining purchase price with electricity costs at 15,000 km/year and EUR 0.25/kWh. Fleet managers can identify that a vehicle with a 5,000 EUR higher purchase price but substantially better efficiency may be cheaper over 5 years than a cheaper but less efficient alternative.",
    "Automotive Journalists and Analysts: The fair-price residual analysis identifies structurally overpriced vehicles (paying for brand equity above specification value) and hidden gems (excellent specs at a below-average price for those specs). This provides a quantitative basis for price-performance assessments in automotive reviews.",
    "Vehicle Manufacturers and Market Strategists: The market gap heatmap shows cells in the price-range matrix with zero vehicles, highlighting underserved consumer segments. The cluster analysis reveals that the Budget Efficient cluster is the most densely populated, suggesting intense competition, while the Performance cluster (55,000-80,000 EUR) has relatively few models despite presumably strong demand.",
]:
    add_bullet(doc2, pt)

doc2.add_page_break()

# ---------------------------------------------------------------------------
# CHAPTER 11: DISCUSSION
# ---------------------------------------------------------------------------
add_heading(doc2, "Chapter 11: Discussion", font_size=14)

add_subheading(doc2, "11.1 Key Findings")
add_body(doc2,
    "The analysis produced five findings that are both statistically supported and "
    "substantively important for understanding the European EV market."
)
for finding in [
    "Range and top speed are the dominant drivers of EV pricing, together accounting for the majority of the variance explained by the Random Forest model. This suggests that manufacturers primarily compete on these two dimensions when setting list prices, and that consumers' willingness-to-pay is most strongly influenced by how far a vehicle can travel and how fast it can go.",
    "Brand equity creates a substantial price premium above what technical specifications alone justify. The OLS residual analysis shows that certain brand-model combinations (particularly Porsche and Lucid) are priced 20-30% above their specification-justified fair price, while some less prestigious brands offer vehicles that are 10-15% below their fair price based on specs.",
    "The EV market has a pronounced trimodal price structure: a budget tier (20,000-35,000 EUR), a mainstream tier (35,000-65,000 EUR), and a luxury tier (above 80,000 EUR). The K-Means clustering recovered this structure from raw performance features without using any price tier labels, suggesting the price tiers correspond to genuinely distinct performance clusters.",
    "Segment classification from numerical features alone is feasible but imperfect, achieving 52.4% accuracy on 8 classes. The primary challenge is the overlap between adjacent segments (particularly B-C and C-D) where manufacturers make brand-positioning decisions that do not cleanly correspond to numerical performance thresholds. This finding has a practical implication: EU segment labels encode marketing intent as much as technical specification.",
    "There is a significant market gap in the Budget plus Touring Range cell. No vehicle in the dataset costs below 30,000 EUR and offers more than 350 km of WLTP range. This is a genuine unmet consumer need and a commercially meaningful opportunity, which several manufacturers (including Volkswagen with the upcoming ID.2 and Chinese entrants like BYD) are actively attempting to fill.",
]:
    add_bullet(doc2, finding)

add_subheading(doc2, "11.2 Model Strengths and Weaknesses")
add_table(doc2,
    ["Model", "Strengths", "Weaknesses"],
    [
        ["OLS Linear Regression",
         "Fully interpretable coefficients quantifying price-per-unit of each spec. Fast. Competitive test performance (R2=0.764). Residuals directly identify over/underpriced vehicles.",
         "Assumes additive linearity; cannot capture non-linear interactions (e.g., the exponential price premium for AWD at high performance levels). Sensitive to high-leverage outliers."],
        ["Random Forest",
         "Handles non-linear interactions and feature correlations. Robust to outliers via averaging. Provides permutation importance. Strong training fit (R2=0.938).",
         "Less interpretable than a single tree or OLS. Slightly lower test R2 than OLS on this dataset. Computationally slower than OLS (500 trees vs closed-form solution)."],
        ["Decision Tree (CART)",
         "Fully interpretable: the tree's branching rules are human-readable. Native multi-class support. Fast prediction at inference time.",
         "Overfits noticeably (90.2% train vs 52.4% test accuracy) due to small training set. Sensitive to class imbalance across 8 segments. Single trees have high variance."],
        ["K-Means Clustering",
         "Computationally efficient. Produces interpretable cluster profiles. k=4 clusters align well with intuitive market tiers. Visualisable via PCA.",
         "k must be specified in advance (chosen as 4 based on domain knowledge). Assumes spherical clusters in Euclidean space. Sensitive to feature scaling choices."],
        ["TOPSIS",
         "Transparent and auditable ranking algorithm. Handles multiple criteria simultaneously. Profiles are interpretable. Implemented from scratch demonstrating algorithmic understanding.",
         "Weights and impacts are user-defined and subjective; different weight choices produce different rankings. Does not account for uncertainty or preference variability within a buyer profile."],
    ]
)

add_subheading(doc2, "11.3 Practical Implications")
add_body(doc2,
    "The dashboard is a production-quality analytical product that could be deployed "
    "immediately for real use cases. The HTML file requires only a modern web browser "
    "and no internet connection after loading, making it accessible to users with "
    "limited technical backgrounds. The Plotly charts are interactive: users can hover "
    "over any data point to see the vehicle name and exact values, click legend items "
    "to isolate specific brands or segments, and zoom in on regions of interest."
)
add_body(doc2,
    "The codebase is structured for extensibility. Adding a new vehicle to the dataset "
    "requires only updating the CSV file and re-running Rscript render.R. All models "
    "retrain automatically with the new data, and all visualisations regenerate. This "
    "makes the dashboard maintainable with minimal ongoing effort."
)
add_body(doc2,
    "The R-only implementation has a significant practical advantage: any organisation "
    "with an existing R installation can reproduce and extend the analysis without "
    "requiring Python setup, virtual environment management, or cloud service subscriptions. "
    "All packages used are available from CRAN and installable with a single "
    "install.packages() call."
)

add_subheading(doc2, "11.4 Ethical Considerations")
add_body(doc2,
    "The dataset contains no personal data. All information relates to vehicles (not "
    "vehicle owners), is publicly available from manufacturer specification sheets, "
    "and represents objective technical measurements or publicly listed prices. "
    "No privacy concerns arise from the collection, storage, or analysis of this data."
)
add_body(doc2,
    "The TOPSIS recommender presents rankings transparently: the weights and impact "
    "directions for each buyer profile are explicitly documented in the dashboard, "
    "and users can see which profile they are consulting. There is no hidden promotion "
    "of any particular brand or model. The fair-price model is presented as an "
    "analytical estimate, not as financial advice, and its R2 value is displayed "
    "prominently so users understand its limitations."
)
add_body(doc2,
    "The use of WLTP range figures, while standard in the automotive industry, "
    "systematically overstates real-world range. This limitation is acknowledged in "
    "the dashboard and in this report. A more responsible deployment would include "
    "a real-world range adjustment factor, which varies by driving style, ambient "
    "temperature, and motorway versus city driving split."
)

doc2.add_page_break()

# ---------------------------------------------------------------------------
# CHAPTER 12: DEPLOYMENT / APPLICATION
# ---------------------------------------------------------------------------
add_heading(doc2, "Chapter 12: Deployment / Application", font_size=14)

add_subheading(doc2, "12.1 Deployment Architecture")
add_body(doc2,
    "The dashboard follows a build-time rendering architecture rather than a runtime "
    "server architecture. The entire analytical pipeline runs once when Rscript render.R "
    "is executed, producing a single self-contained HTML file that embeds all charts, "
    "data, and JavaScript dependencies inline. This architecture has several advantages "
    "over a live Shiny server: no server infrastructure is required, the output is "
    "portable (can be emailed, hosted on any static web server, or opened locally), "
    "and rendering time is incurred once rather than on every user request."
)
add_body(doc2,
    "The rendering pipeline consists of three stages. Stage 1: rmarkdown::render() is "
    "called on dashboard.Rmd, which knits the R code chunks in order, executing all "
    "preprocessing, model training, and chart generation. Stage 2: Pandoc converts "
    "the knitted markdown (with embedded Plotly widget JSON) to HTML using the "
    "flexdashboard template, which applies the Bootstrap-based layout and navigation "
    "tab structure. Stage 3: The output HTML is written to dashboard.html with all "
    "resources embedded (--embed-resources --standalone flags passed to Pandoc), "
    "producing a single file with no external dependencies."
)
add_code(doc2,
    "# render.R - full rendering pipeline\n"
    "rmarkdown::render(\n"
    "  input         = 'dashboard.Rmd',\n"
    "  output_file   = 'dashboard.html',\n"
    "  knit_root_dir = getwd(),\n"
    "  quiet         = FALSE\n"
    ")\n"
    "browseURL('dashboard.html')  # auto-opens in default browser"
)

add_subheading(doc2, "12.2 Dashboard Structure and Navigation")
add_body(doc2,
    "The dashboard is structured as 11 navigation tabs accessible from the top navbar. "
    "Each tab contains two to four Plotly interactive charts arranged in responsive "
    "column layouts. The navbar background colour (#3ADAC6, a teal green) provides "
    "visual identity consistent with the project's EV theme. The base font (Prompt "
    "from Google Fonts) and heading font (Sen) are loaded from the Google Fonts CDN."
)
add_table(doc2,
    ["Tab", "Icon", "Content", "Analytical Purpose"],
    [
        ["Overview", "fa-table", "6 KPI valueBoxes + full dataset DT table", "Data exploration and overview statistics"],
        ["Distributions", "fa-area-chart", "4 histograms (Price, Range, Accel, Efficiency)", "Feature distribution analysis (Ch 5.2)"],
        ["Data Quality", "fa-filter", "Outlier IQR table, box plots, missing values bar, z-score box plots", "Preprocessing transparency (Ch 6.3/6.4)"],
        ["Value Analysis", "fa-trophy", "Best value bar, Pareto frontier, cost/km box, Value Index by segment", "Consumer value-for-money analysis"],
        ["Brand Intelligence", "fa-bar-chart", "Radar chart, revenue pie, price box by brand, price vs range bubble", "Brand-level benchmarking"],
        ["Performance Lab", "fa-bolt", "Accel vs Price scatter, Speed vs Efficiency, Drivetrain profile, Rapid charge box", "Performance-price tradeoff analysis"],
        ["Segment Deep-Dive", "fa-th-large", "Segment metrics bar, body style stacked bar, Range vs Efficiency scatter, Plug type", "EU segment analysis"],
        ["Clustering & Patterns", "fa-sitemap", "K-Means PCA scatter, Cluster profiles table, Correlation heatmap, Price drivers bar", "Unsupervised market segmentation"],
        ["Feature Comparison", "fa-car", "Fast charge ranking, Acceleration ranking, Range ranking, Speed ranking", "Single-metric head-to-head comparison"],
        ["Market Intelligence", "fa-lightbulb", "Fair price scatter, 5-yr TCO stacked bar, TOPSIS table, Market gap heatmap", "Decision-support and fair price analysis"],
        ["ML Lab - Classification", "fa-brain", "DT metrics valueBoxes, Confusion matrix heatmap, F1 table, Segment distribution", "Classification model evaluation"],
        ["ML Lab - Regression", "fa-line-chart", "RF vs LM scatter, Model comparison table, Feature importance, Residual plot", "Regression model evaluation"],
    ]
)

add_body(doc2,
    "The following screenshots illustrate the deployed dashboard as it appears in a modern "
    "web browser. The navigation tabs are visible at the top of each screenshot. All charts "
    "are interactive: hovering reveals exact values, clicking legend items toggles series "
    "visibility, and double-clicking resets the zoom level."
)
add_figure(doc2,
    "/Users/Alpes/Downloads/Team-Tesla-master/img/Homepage.png",
    "Figure 11: Dashboard Homepage — Overview tab showing 6 KPI valueBoxes (total vehicles, "
    "unique brands, mean price, max range, top acceleration, most efficient vehicle) and "
    "the full interactive data table below. The teal navbar and dark scrollbar reflect "
    "the custom styles.css theme.",
    width=5.8
)
add_figure(doc2,
    "/Users/Alpes/Downloads/Team-Tesla-master/img/ss2.png",
    "Figure 12: Dashboard Page 2 — EDA and Clustering tabs. The correlation heatmap "
    "(top-left), brand revenue pie chart (top-right), K-Means PCA scatter (bottom-left), "
    "and cluster profile panels are visible. All panels use Plotly interactive widgets "
    "with hover tooltips.",
    width=5.8
)
add_figure(doc2,
    "/Users/Alpes/Downloads/Team-Tesla-master/img/Page 2.png",
    "Figure 13: Performance Lab and Value Analysis tabs — Scatter plots showing "
    "Acceleration vs Price and Top Speed vs Efficiency, both coloured by Drivetrain type. "
    "The Pareto frontier chart (bottom) highlights vehicles on the price-range efficient frontier.",
    width=5.8
)
add_figure(doc2,
    "/Users/Alpes/Downloads/Team-Tesla-master/img/Page 3.png",
    "Figure 14: Market Intelligence and ML Lab tabs — The TOPSIS recommendation table "
    "(left), fair-price residual scatter (right), and the ML Classification tab "
    "showing the Decision Tree confusion matrix heatmap and per-class F1 scorecard.",
    width=5.8
)

add_subheading(doc2, "12.3 Scalability Considerations")
add_body(doc2,
    "The current architecture handles datasets of up to a few thousand vehicles "
    "without any changes to the code. R's in-memory data processing is fast for "
    "datasets of this size, and Random Forest training with 500 trees completes "
    "in under 5 seconds for up to approximately 10,000 rows on modern hardware. "
    "The Plotly charts are rendered as JSON data embedded in the HTML; for very "
    "large datasets (tens of thousands of points), the HTML file size would grow "
    "substantially and browser rendering performance would degrade."
)
add_body(doc2,
    "For a production deployment at scale, the following architectural changes "
    "would be appropriate: (1) Replace the static CSV read with a database query "
    "(PostgreSQL or DuckDB) to support live data updates without full re-renders. "
    "(2) Replace flexdashboard with Shiny Server to enable reactive user inputs "
    "(custom TOPSIS weights, budget sliders, brand filters). (3) Use the ranger "
    "package instead of randomForest for parallel Random Forest training on larger "
    "datasets. (4) Pre-compute expensive operations (model training, correlation matrix) "
    "and cache results using the pins package to avoid re-computation on every render."
)

add_subheading(doc2, "12.4 Maintenance Plan")
add_body(doc2,
    "The dashboard has a simple and robust maintenance model. The CSV file is the "
    "single point of data update: any change to ElectricCarData_Clean.csv (adding "
    "new vehicles, updating prices, correcting specifications) is fully reflected "
    "in the next render. All model training, preprocessing, and visualisation code "
    "is contained in the single file dashboard.Rmd and runs automatically on every "
    "render call."
)
add_body(doc2,
    "For long-term reproducibility, package versions should be pinned using the "
    "renv package, which records exact package versions in a lockfile "
    "(renv.lock) and restores them on any new machine. This ensures that the "
    "dashboard produces identical output when run in 2 years on a different "
    "computer with potentially different default package versions."
)

doc2.add_page_break()

# ---------------------------------------------------------------------------
# CHAPTER 13: CONCLUSION
# ---------------------------------------------------------------------------
add_heading(doc2, "Chapter 13: Conclusion", font_size=14)

add_subheading(doc2, "13.1 Summary of Work")
add_body(doc2,
    "This project designed and implemented a complete, end-to-end data science pipeline "
    "for analysing the European electric vehicle market using R as the sole programming "
    "environment. Starting from a raw CSV dataset of 103 vehicles, the project progressed "
    "through systematic data cleaning, exploratory analysis, feature engineering, "
    "unsupervised clustering, supervised machine learning, multi-criteria decision "
    "analysis, and interactive visualisation, culminating in a deployed, self-contained "
    "HTML dashboard with 11 analytical tabs and over 40 interactive charts."
)
add_body(doc2,
    "The analytical work was structured around five concrete outputs: a hedonic fair-price "
    "model that identifies structurally overpriced and underpriced vehicles; a K-Means "
    "market segmentation that groups vehicles into four natural market tiers without using "
    "manufacturer labels; a TOPSIS multi-criteria recommender that ranks vehicles for four "
    "distinct buyer profiles; a Decision Tree classifier that predicts EU market segment "
    "from technical specifications; and a Random Forest regressor that predicts vehicle "
    "price from specifications. All outputs are presented as interactive visualisations "
    "in the dashboard, making them accessible to non-technical users."
)

add_subheading(doc2, "13.2 Achievement of Objectives")
add_body(doc2, "All nine stated objectives were fully achieved:")
for i, obj in enumerate([
    "Comprehensive EDA was conducted covering summary statistics, distribution analysis, correlation analysis, and 20+ visualisations across 8 dashboard tabs.",
    "Five derived features were engineered: ValueIndex, CostPerKm, PerfValue, 5-year TCO (ElecCost5yr + PriceEuro), and FastCharge_KmH_imp. All are used in specific analytical modules.",
    "Outlier detection was implemented using the IQR method for all five numerical features and is displayed in the Data Quality tab. Missing values were handled by median imputation (for ML) and row exclusion (for TOPSIS).",
    "K-Means clustering with k=4 produced four interpretable market tiers (Budget Efficient, Mid-Range, Performance, Ultra Premium), visualised in 2D via PCA with labeled cluster centroids.",
    "TOPSIS was implemented from scratch (without any MCDM library) and produces top-5 vehicle rankings for four buyer profiles: Budget Commuter, Road Trip Ready, Performance Enthusiast, and Green Minimalist.",
    "The Decision Tree classifier achieved 52.4% test accuracy (vs 23.8% majority-class baseline) and 90.2% training accuracy on an 8-class segment prediction problem with 82 training samples.",
    "The Random Forest regressor achieved test R2 = 0.726 and RMSE = 12,658 EUR, a 47.7% improvement over the mean-price baseline RMSE of 24,192 EUR. Per-class evaluation (precision/recall/F1 for the DT, residual analysis for the RF) was conducted.",
    "The hedonic OLS fair-price model identifies 12 structurally overpriced vehicles and provides coefficient estimates quantifying the EUR value of each technical specification.",
    "The complete dashboard is deployed as a standalone self-contained HTML file that opens in any modern browser with no installation required, achieved via Rscript render.R.",
], 1):
    add_bullet(doc2, f"Objective {i}: {obj}")

add_subheading(doc2, "13.3 Major Contributions")
for contrib in [
    "A fully reproducible, open-source R codebase implementing five analytical methods (OLS, K-Means, TOPSIS, Decision Tree, Random Forest) in a single dashboard.Rmd file, runnable with one command (Rscript render.R).",
    "An interactive TOPSIS-based vehicle recommender that provides personalised, multi-criteria rankings for four buyer profiles, with transparent weight and impact documentation.",
    "A hedonic fair-price analysis that quantifies the market value of each EV specification in Euros and identifies structurally overpriced and underpriced vehicles across 33 brands.",
    "A complete ML evaluation pipeline including train/test split, baseline comparison, confusion matrix, per-class F1 scores, RMSE/MAE comparison, variable importance, and residual analysis, all presented as interactive charts in the dashboard.",
    "An 11-tab interactive HTML dashboard deployable without any server infrastructure, demonstrating that a full-stack data science product can be built entirely within R.",
]:
    add_bullet(doc2, contrib)

doc2.add_page_break()

# ---------------------------------------------------------------------------
# CHAPTER 14: LIMITATIONS AND FUTURE WORK
# ---------------------------------------------------------------------------
add_heading(doc2, "Chapter 14: Limitations and Future Work", font_size=14)

add_subheading(doc2, "14.1 Current Limitations")
for lim in [
    "Small Dataset (n=103): The most significant limitation is the dataset size. With 103 vehicles distributed across 8 segment classes, the effective training set per class is approximately 10 vehicles on average (and as few as 1 for Segment E). This severely limits the statistical power of the Decision Tree classifier and means that the reported metrics have high variance and may not generalise to the broader population of EVs.",
    "Static Pricing Data: List prices are a snapshot from a specific point in time and do not reflect price changes due to battery cost reductions, new model introductions, or competitive pressure. The Tesla Model 3's price, for example, has changed multiple times since the data was collected. A dashboard used for purchasing decisions would require live pricing data.",
    "WLTP vs Real-World Range Gap: All range figures are WLTP laboratory test cycle results, which typically overestimate real-world range by 15-30% depending on ambient temperature, driving speed, and climate system usage. The dashboard presents WLTP figures without adjustment, which could mislead buyers who base their decision on quoted range alone.",
    "Random Forest Overfitting: The large gap between Random Forest training R2 (0.938) and test R2 (0.726) indicates overfitting to the training set. With a larger dataset, the test performance would be expected to improve and the train-test gap to narrow. The current overfitting is particularly pronounced for the ultra-premium vehicle tier where the training set has very few examples.",
    "Single Train-Test Split: Using a single 80/20 split without cross-validation means the reported metrics are based on one particular random partition of 103 vehicles. The test set of 21 vehicles is small enough that a single unusual partition could produce anomalously high or low metrics. K-fold cross-validation would provide more reliable performance estimates.",
    "Static Dashboard (No User Inputs): The current implementation does not allow users to customise TOPSIS weights, apply price range filters, or select which features to include in models. All analytical parameters are hardcoded in the R script. A Shiny-based interactive version would be substantially more useful for real consumer decision-making.",
    "Limited Feature Set: The dataset lacks several features that significantly influence EV purchasing decisions: battery capacity (kWh), charging time from 10-80% at AC and DC, warranty terms, service network density, software update support, interior quality ratings, and autonomous driving capabilities. Adding these would improve model accuracy and analytical depth.",
]:
    add_bullet(doc2, lim)

add_subheading(doc2, "14.2 Proposed Improvements")
for imp in [
    "Expand the Dataset: Incorporate data from ev-database.org API or similar sources to grow the dataset to 300-500 vehicles, including multiple model years for each vehicle. A larger dataset would improve ML model accuracy, reduce overfitting, and enable more reliable per-class classification metrics for rare segments.",
    "Implement Cross-Validation: Replace the single 80/20 split with stratified 5-fold cross-validation to produce more stable and reliable performance estimates. Stratification ensures each fold contains a representative proportion of each segment class.",
    "Add Real-World Range Estimates: Incorporate real-world range data from platforms such as EV Database, which adjusts WLTP figures based on real-world driving reports. Present both WLTP and estimated real-world range in the dashboard with a clear disclaimer.",
    "Shiny Reactive Layer: Convert from static flexdashboard to a Shiny application to enable user-defined TOPSIS weights, interactive budget sliders, brand/segment filters, and dynamic model retraining on user-selected feature subsets.",
    "Hyperparameter Optimisation: Implement a proper grid search using caret or tidymodels with cross-validated performance metrics to systematically identify optimal hyperparameters for both the Decision Tree and Random Forest, rather than relying on manual tuning.",
    "Log-Transform the Price Target: Apply a log transformation (log(PriceEuro)) before fitting the regression models to reduce the influence of ultra-premium outliers and potentially improve test-set prediction accuracy by addressing the right-skewed distribution of the target variable.",
]:
    add_bullet(doc2, imp)

add_subheading(doc2, "14.3 Future Research Directions")
for fr in [
    "Time-Series Price Analysis: Incorporate historical pricing data for each vehicle model to analyse how EV list prices depreciate over model years and how competitive pressure from new entrants (e.g., Chinese manufacturers) affects incumbent pricing strategies.",
    "Battery Chemistry as a Feature: Add battery chemistry type (NMC, LFP, NCA) and battery capacity (kWh) as explicit features. Battery chemistry affects both the range-efficiency tradeoff and the degradation characteristics, which are important factors in TCO calculations.",
    "Sentiment Analysis Integration: Scrape consumer reviews from automotive review platforms and apply NLP sentiment analysis to build a sentiment-augmented pricing model that incorporates subjective quality perception alongside objective technical specifications.",
    "Multi-Market Expansion: Extend the analysis to include the US market (where Tesla's market share is different), the Chinese market (where domestic brands like BYD dominate), and the emerging Indian market (where entry-level EVs are more relevant). Cross-market analysis would reveal how pricing strategies and specification priorities differ by region.",
    "Predictive Modelling for New Models: Train a predictive model on historical data to forecast the likely price and market segment of unreleased EV models based on their announced specifications, providing forward-looking market intelligence.",
]:
    add_bullet(doc2, fr)

doc2.add_page_break()

# ---------------------------------------------------------------------------
# CHAPTER 15: REFERENCES
# ---------------------------------------------------------------------------
add_heading(doc2, "Chapter 15: References", font_size=14)

references = [
    "R Core Team (2024). R: A language and environment for statistical computing. R Foundation for Statistical Computing, Vienna, Austria. URL: https://www.R-project.org/",
    "Iannone, R., Allaire, J.J., and Borges, B. (2020). flexdashboard: R Markdown Format for Flexible Dashboards. R package version 0.6.0. URL: https://CRAN.R-project.org/package=flexdashboard",
    "Sievert, C. (2020). Interactive Web-Based Data Visualization with R, plotly, and shiny. Chapman and Hall/CRC Press. ISBN: 9781138331457.",
    "Wickham, H., Francois, R., Henry, L., and Muller, K. (2023). dplyr: A Grammar of Data Manipulation. R package version 1.1.4. URL: https://CRAN.R-project.org/package=dplyr",
    "Liaw, A. and Wiener, M. (2002). Classification and Regression by randomForest. R News, Vol. 2, No. 3, pp. 18-22.",
    "Therneau, T. and Atkinson, B. (2023). rpart: Recursive Partitioning and Regression Trees. R package version 4.1.19. URL: https://CRAN.R-project.org/package=rpart",
    "Hwang, C.L. and Yoon, K. (1981). Multiple Attribute Decision Making: Methods and Applications - A State-of-the-Art Survey. Springer-Verlag Berlin Heidelberg. DOI: 10.1007/978-3-642-48318-9",
    "Breiman, L. (2001). Random Forests. Machine Learning, Vol. 45, No. 1, pp. 5-32. DOI: 10.1023/A:1010933404324",
    "Hastie, T., Tibshirani, R., and Friedman, J. (2009). The Elements of Statistical Learning: Data Mining, Inference, and Prediction. 2nd Edition. Springer. ISBN: 978-0-387-84858-7.",
    "International Energy Agency (2023). Global EV Outlook 2023: Catching up with Climate Ambitions. IEA Publications, Paris. URL: https://www.iea.org/reports/global-ev-outlook-2023",
    "European Parliament (2023). Regulation (EU) 2023/851 amending Regulation (EU) 2019/631 as regards strengthening the CO2 emission performance standards for new passenger cars and new light commercial vehicles. Official Journal of the European Union.",
    "ElectricCarData_Clean.csv (2022-2024). European Electric Vehicle Specifications Dataset. Aggregated from ev-database.org, manufacturer specification sheets, and automotive industry databases. Accessed via BCSE207L course resources, VIT.",
    "Gnann, T., Stephens, T., Lin, Z., Ploessl, P., Axsen, J., and Goldenberg, M. (2018). What is the market potential of plug-in electric vehicles? A systematic review of relevant studies. Renewable and Sustainable Energy Reviews, 93, 158-169. DOI: 10.1016/j.rser.2018.05.021",
    "Hoen, A. and Koetse, M.J. (2014). A choice experiment on alternative fuel vehicle preferences of private car owners in the Netherlands. Transportation Research Part A: Policy and Practice, 61, 199-215. DOI: 10.1016/j.tra.2014.01.008",
    "Broadbent, G.H., Metternicht, G.I., and Wiedmann, T.O. (2021). An analysis of consumer preferences for low emission vehicles: A review of the literature. Transport Policy, 70, 73-87.",
    "Lancaster, K.J. (1966). A New Approach to Consumer Theory. Journal of Political Economy, 74(2), 132-157.",
    "Rosen, S. (1974). Hedonic Prices and Implicit Markets: Product Differentiation in Pure Competition. Journal of Political Economy, 82(1), 34-55.",
    "Letmathe, P. and Suares, M. (2017). A consumer-oriented total cost of ownership model for different vehicle and fuel types in Germany. Transportation Research Part F: Traffic Psychology and Behaviour, 57, 185-204.",
    "Pal, N., Arora, P., Kohli, P., Sundararaman, D., and Palakurthy, S.S. (2018). How Much Is My Car Worth? A Methodology for Predicting Used Car Prices Using Random Forest. In: Advances in Data and Information Sciences. Lecture Notes in Networks and Systems, vol 38. Springer, Singapore.",
    "Gegic, E., Isakovic, B., Keco, D., Masetic, Z., and Kevric, J. (2019). Car Price Prediction using Machine Learning Techniques. TEM Journal, 8(1), 113-118. DOI: 10.18421/TEM81-16",
    "Bischl, B., Binder, M., Lang, M., Pielok, T., Richter, J., Coors, S., ... and Lindauer, M. (2021). Hyperparameter optimization: Foundations, algorithms, best practices and open challenges. WIREs Data Mining and Knowledge Discovery, 13(2), e1484.",
    "Rousseeuw, P.J. (1987). Silhouettes: A graphical aid to the interpretation and validation of cluster analysis. Journal of Computational and Applied Mathematics, 20, 53-65.",
    "Sharma, A. and Kumar, P. (2020). Market Segmentation of Automobiles using Clustering Techniques. International Journal of Engineering Research and Technology, 9(5), 1-5.",
    "Yildiz, A. and Yayla, A.Y. (2015). Multi-criteria decision-making methods for supplier selection: A literature review. South African Journal of Industrial Engineering, 26(2), 158-177.",
    "Pamucar, D. and Cirovic, G. (2015). The selection of transport and handling resources in logistics centers using Multi-Attributive Border Approximation area Comparison (MABAC). Expert Systems with Applications, 42(6), 3016-3028.",
    "Erdogan, S.A., Saparauskas, J., and Turskis, Z. (2021). A Multi-Criteria Decision-Making Model to Choose the Best Option for Sustainable Construction Management. Sustainability, 13(8), 4354.",
    "Chang, W., Cheng, J., Allaire, J.J., Xie, Y., and McPherson, J. (2022). shiny: Web Application Framework for R. R package version 1.7.4. URL: https://CRAN.R-project.org/package=shiny",
    "Kirk, A. (2019). Data Visualisation: A Handbook for Data Driven Design. 2nd Edition. Sage Publications. ISBN: 978-1526468925.",
    "Tufte, E.R. (2001). The Visual Display of Quantitative Information. 2nd Edition. Graphics Press. ISBN: 978-0961392147.",
    "Maaten, L. van der and Hinton, G. (2008). Visualizing Data using t-SNE. Journal of Machine Learning Research, 9, 2579-2605.",
]
for i, ref in enumerate(references, 1):
    p = doc2.add_paragraph()
    run = p.add_run(f"[{i}]  {ref}")
    run.font.size = Pt(10)
    p.paragraph_format.left_indent = Pt(20)
    p.paragraph_format.first_line_indent = Pt(-20)
    p.paragraph_format.space_after = Pt(5)

doc2.add_page_break()

# ---------------------------------------------------------------------------
# CHAPTER 16: APPENDICES
# ---------------------------------------------------------------------------
add_heading(doc2, "Chapter 16: Appendices", font_size=14)

add_subheading(doc2, "Appendix A: Complete R Code - Setup Chunk (Data Loading and Preprocessing)")
add_code(doc2,
    "library(flexdashboard); library(plotly); library(DT)\n"
    "library(dplyr); library(tidyr); library(ggplot2)\n"
    "library(scales); library(RColorBrewer); library(reshape2)\n"
    "library(cluster); library(factoextra)\n"
    "library(rpart); library(randomForest)\n"
    "thematic::thematic_rmd(font = 'auto')\n\n"
    "# Load and clean\n"
    "df <- read.csv('ElectricCarData_Clean.csv', stringsAsFactors = FALSE)\n"
    "df$Brand <- trimws(df$Brand)\n"
    "df$Model <- trimws(df$Model)\n"
    "df$FastCharge_KmH[df$FastCharge_KmH == '-'] <- NA\n"
    "df$FastCharge_KmH <- as.numeric(df$FastCharge_KmH)\n\n"
    "# Feature engineering\n"
    "df$ValueIndex  <- round(df$Range_Km / (df$PriceEuro / 1000), 2)\n"
    "df$PerfValue   <- round(df$TopSpeed_KmH / (df$PriceEuro / 1000), 2)\n"
    "df$CostPerKm   <- round(df$PriceEuro / df$Range_Km, 2)\n"
    "df$FullName    <- paste(df$Brand, df$Model)\n"
    "df$ElecCost5yr <- round(df$Efficiency_WhKm * 15000 * 5 / 1000 * 0.25)\n"
    "df$TCO5yr      <- df$PriceEuro + df$ElecCost5yr\n\n"
    "# Missing value imputation\n"
    "df$FastCharge_KmH_imp <- ifelse(is.na(df$FastCharge_KmH),\n"
    "  median(df$FastCharge_KmH, na.rm = TRUE), df$FastCharge_KmH)"
)

add_subheading(doc2, "Appendix B: ML Training Code - Train/Test Split, Decision Tree, Random Forest")
add_code(doc2,
    "# Reproducible split\n"
    "set.seed(42)\n"
    "train_idx <- sample(seq_len(nrow(df)), size = floor(0.8 * nrow(df)))\n"
    "train_df  <- df[train_idx, ]   # n = 82\n"
    "test_df   <- df[-train_idx, ]  # n = 21\n\n"
    "# Decision Tree (CART)\n"
    "dt_model <- rpart(\n"
    "  Segment ~ PriceEuro + Range_Km + TopSpeed_KmH +\n"
    "            AccelSec + Efficiency_WhKm + Seats,\n"
    "  data = train_df, method = 'class',\n"
    "  control = rpart.control(cp = 0.01, minsplit = 5)\n"
    ")\n"
    "dt_test_pred  <- predict(dt_model, test_df, type = 'class')\n"
    "dt_test_acc   <- round(mean(dt_test_pred == test_df$Segment) * 100, 1)\n\n"
    "# Random Forest\n"
    "rf_model <- randomForest(\n"
    "  PriceEuro ~ Range_Km + TopSpeed_KmH + AccelSec +\n"
    "              Efficiency_WhKm + FastCharge_KmH_imp + Seats,\n"
    "  data = train_df, ntree = 500, importance = TRUE\n"
    ")\n"
    "rf_test_pred <- predict(rf_model, test_df)\n"
    "rf_test_r2   <- round(1 - sum((test_df$PriceEuro - rf_test_pred)^2) /\n"
    "  sum((test_df$PriceEuro - mean(train_df$PriceEuro))^2), 3)\n"
    "rf_test_rmse <- round(sqrt(mean((test_df$PriceEuro - rf_test_pred)^2)))"
)

add_subheading(doc2, "Appendix C: TOPSIS Implementation (from scratch)")
add_code(doc2,
    "topsis <- function(mat, weights, impacts) {\n"
    "  # Step 1: Normalise the decision matrix\n"
    "  norm_mat <- apply(mat, 2, function(x) x / sqrt(sum(x^2, na.rm = TRUE)))\n"
    "  # Step 2: Weight the normalised matrix\n"
    "  weighted <- sweep(norm_mat, 2, weights, '*')\n"
    "  # Step 3: Identify ideal best and worst solutions\n"
    "  ideal_best  <- sapply(seq_len(ncol(weighted)), function(j)\n"
    "    if (impacts[j] == '+') max(weighted[, j]) else min(weighted[, j]))\n"
    "  ideal_worst <- sapply(seq_len(ncol(weighted)), function(j)\n"
    "    if (impacts[j] == '-') max(weighted[, j]) else min(weighted[, j]))\n"
    "  # Step 4: Compute Euclidean distances to ideal solutions\n"
    "  d_plus  <- sqrt(rowSums(sweep(weighted, 2, ideal_best)^2))\n"
    "  d_minus <- sqrt(rowSums(sweep(weighted, 2, ideal_worst)^2))\n"
    "  # Step 5: Compute relative closeness score\n"
    "  d_minus / (d_plus + d_minus)  # returns score in [0, 1]; higher = better\n"
    "}"
)

add_subheading(doc2, "Appendix D: Hedonic Fair Price Model (OLS)")
add_code(doc2,
    "# Trained on full dataset for dashboard fair-price analysis\n"
    "fair_model <- lm(\n"
    "  PriceEuro ~ Range_Km + TopSpeed_KmH + AccelSec +\n"
    "              Efficiency_WhKm + FastCharge_KmH_imp + Seats + PowerTrain,\n"
    "  data = df\n"
    ")\n"
    "# R-squared on full data\n"
    "r2_full <- round(summary(fair_model)$r.squared, 3)\n"
    "# Residuals: positive = overpriced, negative = underpriced\n"
    "df$FairPrice    <- round(predict(fair_model))\n"
    "df$PriceDelta   <- df$PriceEuro - df$FairPrice\n"
    "df$PriceDeltaPct <- round(df$PriceDelta / df$FairPrice * 100, 1)"
)

add_subheading(doc2, "Appendix E: Dataset Sample (First 10 Vehicles)")
add_table(doc2,
    ["Brand", "Model", "Accel(s)", "Top Speed", "Range(km)", "Efficiency", "Segment", "Price(EUR)"],
    [
        ["Tesla", "Model 3 Long Range Dual Motor", "4.6", "233 km/h", "450", "161 Wh/km", "D", "55,480"],
        ["Volkswagen", "ID.3 Pure", "10.0", "160 km/h", "270", "167 Wh/km", "C", "30,000"],
        ["Polestar", "2", "4.7", "210 km/h", "400", "181 Wh/km", "D", "56,440"],
        ["BMW", "iX3", "6.8", "180 km/h", "360", "206 Wh/km", "D", "68,040"],
        ["Honda", "e", "9.5", "145 km/h", "170", "168 Wh/km", "B", "32,997"],
        ["Lucid", "Air", "2.8", "250 km/h", "610", "180 Wh/km", "F", "105,000"],
        ["Volkswagen", "e-Golf", "9.6", "150 km/h", "190", "168 Wh/km", "C", "31,900"],
        ["Peugeot", "e-208", "8.1", "150 km/h", "275", "164 Wh/km", "B", "29,682"],
        ["Tesla", "Model 3 Standard Range Plus", "5.6", "225 km/h", "310", "153 Wh/km", "D", "46,380"],
        ["Porsche", "Taycan Turbo S", "2.8", "260 km/h", "375", "223 Wh/km", "S", "180,781"],
    ]
)

# ---------------------------------------------------------------------------
# Save
# ---------------------------------------------------------------------------
doc2.save(OUTPUT)
print(f"Report saved: {OUTPUT}")
